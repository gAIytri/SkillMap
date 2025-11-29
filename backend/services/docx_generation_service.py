"""
DOCX Generation Service - Pure python-docx Programmatic Approach
Generates resume DOCX from JSON data with full styling control
No templates, no Jinja2 - just clean, reliable document building
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
from typing import Dict, Any, List
import logging

logger = logging.getLogger(__name__)

# ============================================================================
# STYLE CONFIGURATION (matching your resume exactly)
# ============================================================================

# Global spacing configuration (in Pt)
SECTION_SPACING_BEFORE = 5  # Space before each section header (reduced for tighter spacing)
SECTION_SPACING_AFTER = 2   # Space after each section header (NO gap after underline)

STYLES = {
    'contact': {
        'style': 'Normal',
        'alignment': WD_ALIGN_PARAGRAPH.CENTER,
        'font_size': 10,
    },
    'section_header': {
        'style': 'Heading 1',
        'underline': True,
    },
    'job_header': {
        'style': 'Heading 2',
    },
    'bullet': {
        'style': 'List Paragraph',
        'font_size': 10,
    },
    'normal': {
        'style': 'Normal',
        'font_size': 10,
    },
    'education_school': {
        'style': 'Normal',
        'font_size': 10,
        'bold': True,
        'alignment': WD_ALIGN_PARAGRAPH.JUSTIFY,
    },
    'education_degree': {
        'style': 'Normal',
        'font_size': 10,
        'italic': True,
        'alignment': WD_ALIGN_PARAGRAPH.LEFT,
    },
}


# ============================================================================
# HELPER FUNCTIONS
# ============================================================================

def sanitize_text(text: str) -> str:
    """
    Remove invalid XML characters and normalize special Unicode characters

    Args:
        text: Input text that may contain invalid characters

    Returns:
        Sanitized text safe for XML/DOCX with normalized characters
    """
    if not text:
        return ""

    import re
    import html

    # Decode HTML entities first (like &amp;, &lt;, &gt;, etc.)
    text = html.unescape(text)

    # Normalize common Unicode characters to their ASCII equivalents
    replacements = {
        # Quotes (smart quotes to regular quotes)
        '\u2018': "'",  # Left single quote
        '\u2019': "'",  # Right single quote
        '\u201C': '"',  # Left double quote
        '\u201D': '"',  # Right double quote
        '\u2033': '"',  # Double prime
        '\u2032': "'",  # Prime

        # Dashes (en-dash, em-dash to hyphen)
        '\u2013': '-',  # En dash
        '\u2014': '-',  # Em dash
        '\u2212': '-',  # Minus sign

        # Spaces
        '\u00A0': ' ',  # Non-breaking space
        '\u2009': ' ',  # Thin space
        '\u200B': '',   # Zero-width space

        # Ellipsis
        '\u2026': '...',  # Horizontal ellipsis

        # Ampersand variants (fullwidth and small variants)
        '\uFF06': '&',  # Fullwidth ampersand ＆
        '\uFE60': '&',  # Small ampersand ﹠
    }

    # Apply replacements
    for old_char, new_char in replacements.items():
        text = text.replace(old_char, new_char)

    # Remove null bytes and control characters (except newline, carriage return, tab)
    # Keep only valid XML characters
    # Valid: \x09 (tab), \x0A (LF), \x0D (CR), \x20-\uD7FF, \uE000-\uFFFD
    text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F]', '', text)

    return text


def add_hyperlink(paragraph, text: str, url: str, size: int = 10, color: RGBColor = None):
    """
    Add a clickable hyperlink to a paragraph

    Args:
        paragraph: Paragraph object
        text: Display text
        url: URL to link to
        size: Font size in points
        color: Font color (RGBColor object)
    """
    # This gets access to the document.xml.rels file
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    # Create the hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a new run for the hyperlink text
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Set hyperlink style (blue color, no underline)
    if color:
        c = OxmlElement('w:color')
        # RGBColor is a tuple (r, g, b), not an object with .r, .g, .b
        c.set(qn('w:val'), '%02x%02x%02x' % (color[0], color[1], color[2]))
        rPr.append(c)

    # Set font size
    if size:
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(size * 2))  # Word uses half-points
        rPr.append(sz)

    # Set font name to Calibri
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Calibri')
    rFonts.set(qn('w:hAnsi'), 'Calibri')
    rPr.append(rFonts)

    new_run.append(rPr)

    # Add text
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink


def add_bullet_paragraph(doc: Document, text: str, font_size: int = 10, keep_together: bool = True):
    """
    Add a bulleted paragraph with visible bullet points (•)

    Args:
        doc: Document object
        text: Bullet text
        font_size: Font size in points (default 10)
        keep_together: If True, prevents bullet from splitting across pages (default True)
    """
    # Create paragraph and add bullet character directly
    para = doc.add_paragraph()

    # Add bullet character (•) followed by text
    run = para.add_run(f"• {sanitize_text(text)}")
    run.font.size = Pt(font_size)
    run.font.name = 'Calibri'

    # Set hanging indent for bullet alignment
    para.paragraph_format.left_indent = Inches(0.12)
    para.paragraph_format.first_line_indent = Inches(-0.10)

    # Set JUSTIFY alignment for bullets
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Ensure no extra spacing
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0  # Single line spacing

    # Pagination: Keep bullet paragraph together (don't split mid-bullet)
    if keep_together:
        para.paragraph_format.keep_together = True

    return para


def add_header_section(doc: Document, personal_info: Dict[str, Any]):
    """
    Add header section with name, role, and contact info

    Format:
    NAME (centered, 18pt, bold)

    If 1-2 links:
        Current Role | location | email | phone | LinkedIn | GitHub (centered, 10pt)

    If 3+ links:
        Current Role | location | email | phone (centered, 10pt)
        LinkedIn | GitHub | Portfolio (centered, 10pt, blue links)

    Uses paragraphs with EXACTLY spacing for precise control.
    """
    # Paragraph 1: Name (18pt, bold)
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(sanitize_text(personal_info.get('name', '')))
    name_run.font.size = Pt(18)
    name_run.font.bold = True
    name_run.font.name = 'Calibri'
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Use EXACTLY spacing for name line
    name_para.paragraph_format.space_before = Pt(0)
    name_para.paragraph_format.space_after = Pt(2)  # Minimal gap after name
    name_para.paragraph_format.line_spacing = Pt(18)  # Exact 18pt for 18pt font
    name_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY

    # Paragraph 2: Contact info (10pt) - with smart link placement
    contact_parts = []

    # Add current_role FIRST if present (check for non-empty string)
    if personal_info.get('current_role') and personal_info['current_role'].strip():
        contact_parts.append(personal_info['current_role'].strip())

    if personal_info.get('location') and personal_info['location'].strip():
        contact_parts.append(personal_info['location'].strip())
    if personal_info.get('email') and personal_info['email'].strip():
        contact_parts.append(personal_info['email'].strip())
    if personal_info.get('phone') and personal_info['phone'].strip():
        contact_parts.append(personal_info['phone'].strip())

    header_links = personal_info.get('header_links', [])

    # Determine if links should be on same line or separate line
    # If 1-2 links: add to same line
    # If 3+ links: put on separate line
    links_on_same_line = len(header_links) <= 2

    if contact_parts or header_links:
        contact_links_para = doc.add_paragraph()
        contact_links_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Use EXACTLY 11pt spacing for tight 10pt text lines
        contact_links_para.paragraph_format.space_before = Pt(0)
        contact_links_para.paragraph_format.space_after = Pt(4)
        contact_links_para.paragraph_format.line_spacing = Pt(11)  # Tight for 10pt font
        contact_links_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY

        # Add contact info
        if contact_parts:
            contact_run = contact_links_para.add_run(' | '.join(contact_parts))
            contact_run.font.size = Pt(10)
            contact_run.font.name = 'Calibri'

        # Add links on same line (if 1-2 links) or next line (if 3+ links)
        if header_links:
            # If 3+ links, add line break before links
            if not links_on_same_line:
                contact_links_para.add_run('\n')

            for idx, link in enumerate(header_links):
                text = link.get('text', '')
                url = link.get('url', None)

                if not text:
                    continue

                # Add separator before link (including first link if on same line with contact info)
                if idx > 0 or (links_on_same_line and contact_parts):
                    separator_run = contact_links_para.add_run(' | ')
                    separator_run.font.size = Pt(10)
                    separator_run.font.name = 'Calibri'

                # Create hyperlink if URL exists, otherwise plain text
                if url:
                    add_hyperlink(contact_links_para, text, url, size=10, color=RGBColor(0, 0, 0))  # Black color
                else:
                    link_run = contact_links_para.add_run(text)
                    link_run.font.size = Pt(10)
                    link_run.font.name = 'Calibri'
                    link_run.font.color.rgb = RGBColor(0, 0, 0)  # Black color


def add_section_header(doc: Document, title: str, keep_with_next: bool = False):
    """
    Add section header with underline (e.g., PROFESSIONAL SUMMARY)

    Args:
        doc: Document object
        title: Section title text
        keep_with_next: If True, keeps header on same page as next paragraph (for small sections)
    """
    # Create plain paragraph WITHOUT Heading1 style to avoid built-in spacing
    # All formatting (bold, size, font, underline) is applied manually below
    para = doc.add_paragraph()

    # Add title text as a run
    run = para.add_run(title)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = 'Calibri'

    # Ensure header starts at content area edge (no additional indent)
    para.paragraph_format.left_indent = Inches(0)
    para.paragraph_format.first_line_indent = Inches(0)

    # Add spacing before section, NO spacing after
    para.paragraph_format.space_before = Pt(SECTION_SPACING_BEFORE)
    para.paragraph_format.space_after = Pt(SECTION_SPACING_AFTER)  # NO GAP after underline
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE  # Force exact single spacing

    # Add bottom border (underline effect for the entire paragraph)
    # The border will extend full width of the content area (between margins)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')  # Border thickness
    bottom.set(qn('w:space'), '0')  # No spacing - border goes edge to edge
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Force paragraph to span full width by setting indentation explicitly
    # This ensures the border extends edge-to-edge within margins
    pPr_ind = para._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '0')
    ind.set(qn('w:right'), '0')
    pPr_ind.append(ind)

    # Pagination control: Keep header with next paragraph if requested
    if keep_with_next:
        para.paragraph_format.keep_with_next = True

    return para


def add_professional_summary(doc: Document, summary: str, section_name: str = 'PROFESSIONAL SUMMARY'):
    """Add professional summary section"""
    if not summary:
        return

    # Keep header with content (entire section should stay together)
    add_section_header(doc, section_name, keep_with_next=True)

    summary_para = doc.add_paragraph(sanitize_text(summary))
    summary_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # JUSTIFY aligned (both left and right edges)
    summary_para.paragraph_format.left_indent = Inches(0)
    summary_para.paragraph_format.first_line_indent = Inches(0)
    summary_para.paragraph_format.space_before = Pt(0)  # NO gap after underline
    summary_para.paragraph_format.space_after = Pt(0)
    summary_para.paragraph_format.line_spacing = 1.0  # Single line spacing
    # Pagination: Keep summary together (don't split across pages)
    summary_para.paragraph_format.keep_together = True
    if summary_para.runs:
        summary_para.runs[0].font.size = Pt(10)
        summary_para.runs[0].font.name = 'Calibri'


def add_education_section(doc: Document, education: List[Dict[str, Any]], section_name: str = 'EDUCATION'):
    """
    Add education section

    Format:
    EDUCATION (Heading 1, underlined)
    Institution, Location (Normal, 10pt, bold)
    Degree – GPA: X.X/Y (optional /Y)    Date (Normal, 10pt, italic, right-aligned)

    Pagination: Entire section stays together (education is typically short)
    """
    if not education:
        return

    # Keep header with content (entire section should stay together)
    add_section_header(doc, section_name, keep_with_next=True)

    for edu in education:
        # School name and location (bold) with date on same line
        school_parts = [edu.get('institution', '')]
        if edu.get('location'):
            school_parts.append(edu['location'])

        school_para = doc.add_paragraph(style='Normal')
        school_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Explicitly set indents to 0 to align with section header
        school_para.paragraph_format.left_indent = Pt(1)
        school_para.paragraph_format.first_line_indent = Inches(0)
        school_para.paragraph_format.space_before = Pt(0)
        school_para.paragraph_format.space_after = Pt(0)
        school_para.paragraph_format.line_spacing = 1.0  # Single line spacing
        # Pagination: Keep school with degree line
        school_para.paragraph_format.keep_with_next = True
        school_para.paragraph_format.keep_together = True

        # Add school name and location (bold)
        school_run = school_para.add_run(', '.join(school_parts))
        school_run.font.bold = True
        school_run.font.size = Pt(10)

        # Add date on right side with tab (SAME line as school, at underline end)
        if edu.get('graduation_date'):
            tab_stops = school_para.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(7.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT)

            # Add tab and date (bold to match original resume)
            date_run = school_para.add_run(f"\t{edu['graduation_date']}")
            date_run.font.bold = True
            date_run.font.size = Pt(10)

        # Degree line (italic) - NO date on this line
        degree_parts = [edu.get('degree', '')]
        if edu.get('gpa'):
            gpa_text = edu['gpa']
            # Add "out of" if provided (e.g., "3.5/4")
            if edu.get('gpa_out_of'):
                gpa_text = f"{gpa_text}/{edu['gpa_out_of']}"
            degree_parts.append(f"– GPA: {gpa_text}")

        degree_para = doc.add_paragraph(style='Normal')
        degree_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Explicitly set indents to 0 to align with section header
        degree_para.paragraph_format.left_indent = Inches(0)
        degree_para.paragraph_format.first_line_indent = Inches(0)
        degree_para.paragraph_format.space_before = Pt(0)
        degree_para.paragraph_format.space_after = Pt(0)
        degree_para.paragraph_format.line_spacing = 1.0  # Single line spacing
        # Pagination: Keep degree paragraph together
        degree_para.paragraph_format.keep_together = True

        # Add degree text (italic, no date)
        degree_run = degree_para.add_run(' '.join(degree_parts))
        degree_run.font.italic = True
        degree_run.font.size = Pt(10)


def add_experience_section(doc: Document, experience: List[Dict[str, Any]], section_name: str = 'EXPERIENCE'):
    """
    Add experience section

    Format:
    EXPERIENCE (Heading 1, underlined)
    Company Name                         Start – End (bold, 11pt, date right-aligned)
    Role, Location (normal weight, 10pt, black, closer spacing)
    • Bullet 1 (List Paragraph, 10pt)
    • Bullet 2

    Pagination: Each individual job entry stays together (header + bullets),
    but section can span multiple pages
    """
    if not experience:
        return

    # Don't use keep_with_next here - allow section to span pages
    add_section_header(doc, section_name, keep_with_next=False)

    for idx, exp in enumerate(experience):
        # Job header (Company name + dates)
        company = exp.get('company', '')
        title = exp.get('title', '')
        location = exp.get('location', '')
        start_date = exp.get('start_date', '')
        end_date = exp.get('end_date', '')

        # Line 1: Company Name (bold) + Dates (right-aligned, bold)
        company_para = doc.add_paragraph()

        # Explicitly set indents to 0 to align with section header
        company_para.paragraph_format.left_indent = Pt(1)
        company_para.paragraph_format.first_line_indent = Inches(0)
        # First item: minimal gap after section underline. Subsequent items: add spacing
        company_para.paragraph_format.space_before = Pt(2) if idx == 0 else Pt(6)
        company_para.paragraph_format.space_after = Pt(0)  # NO gap after company line
        company_para.paragraph_format.line_spacing = Pt(11)  # Tight exact spacing for 11pt font
        company_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        # Pagination: Keep company with role line
        company_para.paragraph_format.keep_with_next = True
        company_para.paragraph_format.keep_together = True

        if company:
            # Add company name (bold, 11pt)
            company_run = company_para.add_run(company)
            company_run.font.bold = True
            company_run.font.size = Pt(11)
            company_run.font.name = 'Calibri'

            # Add dates on right side (bold, 11pt)
            if start_date and end_date:
                tab_stops = company_para.paragraph_format.tab_stops
                tab_stops.add_tab_stop(Inches(7.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT)

                date_run = company_para.add_run(f"\t{start_date} – {end_date}")
                date_run.font.bold = True
                date_run.font.size = Pt(11)
                date_run.font.name = 'Calibri'

        # Line 2: Role, Location (light weight, 10pt, very close to company line)
        role_para = doc.add_paragraph()

        # Explicitly set indents to 0 to align with section header
        role_para.paragraph_format.left_indent = Pt(1)
        role_para.paragraph_format.first_line_indent = Inches(0)
        # EXTREMELY tight spacing - minimal gap between company and role line
        role_para.paragraph_format.space_before = Pt(0)  # Zero spacing to keep very close to company above
        role_para.paragraph_format.space_after = Pt(2)  # Small gap after role line before bullets
        role_para.paragraph_format.line_spacing = Pt(10)  # Exact line height for 10pt font
        role_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        # Pagination: Keep role with bullets
        role_para.paragraph_format.keep_with_next = True
        role_para.paragraph_format.keep_together = True

        # Build role and location text
        role_parts = []
        if title:
            role_parts.append(title)
        if location:
            role_parts.append(location)

        if role_parts:
            role_text = ', '.join(role_parts)
            role_run = role_para.add_run(role_text)
            role_run.font.bold = False  # Normal weight (not bold)
            role_run.font.size = Pt(10)
            role_run.font.name = 'Calibri'
            role_run.font.color.rgb = RGBColor(0, 0, 0)  # Normal black color (not gray)

        # Bullets (List Paragraph, 10pt)
        bullets = exp.get('bullets', [])
        for i, bullet in enumerate(bullets):
            # Keep all bullets together with the job header
            # Set keep_with_next=True on all bullets except the last one
            is_last_bullet = (i == len(bullets) - 1)
            is_first_bullet = (i == 0)

            bullet_para = add_bullet_paragraph(doc, bullet, font_size=10, keep_together=True)

            # Add extra space before FIRST bullet only (after role line)
            if is_first_bullet:
                bullet_para.paragraph_format.space_before = Pt(3)

            if not is_last_bullet:
                bullet_para.paragraph_format.keep_with_next = True


def add_projects_section(doc: Document, projects: List[Dict[str, Any]], section_name: str = 'PROJECTS'):
    """
    Add projects section

    Format:
    PROJECTS (Heading 1, underlined)
    Project Name – Technologies    Date (Heading 2)
    • Description bullet 1 (List Paragraph, 10pt)
    • Description bullet 2

    Pagination: Each individual project entry stays together (header + bullets),
    but section can span multiple pages
    """
    if not projects:
        return

    # Don't use keep_with_next here - allow section to span pages
    add_section_header(doc, section_name, keep_with_next=False)

    for idx, project in enumerate(projects):
        # Project header (similar to experience: name + date on line 1, technologies on line 2)
        name = project.get('name', '')
        technologies = project.get('technologies', [])
        link = project.get('link', '')
        date = project.get('date', '')

        # Build tech string first to determine spacing
        tech_str = ''
        if technologies:
            if isinstance(technologies, list):
                tech_str = ', '.join(technologies)
            else:
                tech_str = str(technologies)

        # Line 1: Project Name (bold) + Date (right-aligned)
        project_para = doc.add_paragraph()

        # Explicitly set indents to 0 to align with section header
        project_para.paragraph_format.left_indent = Pt(1)
        project_para.paragraph_format.first_line_indent = Inches(0)
        # First item: minimal gap after section underline. Subsequent items: add spacing
        project_para.paragraph_format.space_before = Pt(2) if idx == 0 else Pt(6)
        # If no technologies, add MORE spacing after project name (bullets have 0 space_before). Otherwise, no gap (tech line will be tight)
        project_para.paragraph_format.space_after = Pt(2) if not tech_str else Pt(0)
        project_para.paragraph_format.line_spacing = Pt(11)  # Tight exact spacing for 11pt font
        project_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        # Pagination: Keep project name with next element (either tech line or bullets)
        project_para.paragraph_format.keep_with_next = True
        project_para.paragraph_format.keep_together = True

        if name:
            # Add project name (bold, 11pt)
            project_run = project_para.add_run(sanitize_text(name))
            project_run.font.bold = True
            project_run.font.size = Pt(11)
            project_run.font.name = 'Calibri'

            # Add date on right side (bold, 11pt)
            if date:
                tab_stops = project_para.paragraph_format.tab_stops
                tab_stops.add_tab_stop(Inches(7.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT)

                date_run = project_para.add_run(f"\t{date}")
                date_run.font.bold = True
                date_run.font.size = Pt(11)
                date_run.font.name = 'Calibri'

        # Line 2: Technologies (light weight, 10pt, very close to project name line)
        if tech_str:
            tech_para = doc.add_paragraph()

            # Explicitly set indents to 0 to align with section header
            tech_para.paragraph_format.left_indent = Pt(1)
            tech_para.paragraph_format.first_line_indent = Inches(0)
            # EXTREMELY tight spacing - minimal gap between project name and tech line
            tech_para.paragraph_format.space_before = Pt(0)  # Zero spacing to keep very close to project name above
            tech_para.paragraph_format.space_after = Pt(2)  # Small gap after tech line before bullets
            tech_para.paragraph_format.line_spacing = Pt(10)  # Exact line height for 10pt font
            tech_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            # Pagination: Keep tech with bullets
            tech_para.paragraph_format.keep_with_next = True
            tech_para.paragraph_format.keep_together = True

            # Add technologies text (normal weight, 10pt)
            tech_run = tech_para.add_run(sanitize_text(tech_str))
            tech_run.font.bold = False
            tech_run.font.size = Pt(10)
            tech_run.font.name = 'Calibri'

        # Get bullets from either bullets array or description string (backward compatibility)
        bullets = project.get('bullets', [])
        description = project.get('description', '')

        # Collect all bullet points
        all_bullets = []
        if bullets and isinstance(bullets, list) and len(bullets) > 0:
            # Use bullets array directly (new format)
            all_bullets = [str(b).strip() for b in bullets if str(b).strip()]
        elif description and isinstance(description, str):
            # Fall back to description string (old format) - split by newlines
            all_bullets = [line.strip() for line in description.split('\n') if line.strip()]

        # Add bullets with keep_with_next to keep project entry together
        for i, bullet in enumerate(all_bullets):
            is_first_bullet = (i == 0)
            is_last_bullet = (i == len(all_bullets) - 1)
            bullet_para = add_bullet_paragraph(doc, bullet, font_size=10, keep_together=True)

            # If this is the first bullet AND there are no technologies, add spacing before it
            if is_first_bullet and not tech_str:
                bullet_para.paragraph_format.space_before = Pt(2)

            if not is_last_bullet:
                bullet_para.paragraph_format.keep_with_next = True


def add_skills_section(doc: Document, skills: List[Dict[str, Any]], section_name: str = 'TECHNICAL SKILLS'):
    """
    Add technical skills section

    Format:
    TECHNICAL SKILLS (Heading 1, underlined)
    Category: skill1, skill2, skill3 (Normal, 10pt, bold category)

    Pagination: Entire section stays together (skills are typically short)
    """
    if not skills:
        return

    # Keep header with content (entire section should stay together)
    add_section_header(doc, section_name, keep_with_next=True)

    for skill_cat in skills:
        category = skill_cat.get('category', '')
        skill_items = skill_cat.get('skills', [])

        if category and skill_items:
            skill_para = doc.add_paragraph(style='Normal')
            skill_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # JUSTIFY alignment
            # Explicitly set indents to 0 to align with section header
            skill_para.paragraph_format.left_indent = Pt(1)
            skill_para.paragraph_format.first_line_indent = Inches(0)
            skill_para.paragraph_format.space_before = Pt(0)
            skill_para.paragraph_format.space_after = Pt(0)
            skill_para.paragraph_format.line_spacing = 1.0  # Single line spacing
            # Pagination: Keep skill lines together
            skill_para.paragraph_format.keep_with_next = True
            skill_para.paragraph_format.keep_together = True

            # Category (bold)
            cat_run = skill_para.add_run(f"{category}: ")
            cat_run.font.bold = True
            cat_run.font.size = Pt(10)

            # Skills (normal)
            skills_text = ', '.join(skill_items) if isinstance(skill_items, list) else str(skill_items)
            items_run = skill_para.add_run(skills_text)
            items_run.font.size = Pt(10)


def add_certifications_section(doc: Document, certifications: List[str], section_name: str = 'CERTIFICATIONS'):
    """
    Add certifications section

    Format:
    CERTIFICATIONS (Heading 1, underlined)
    • Certification 1 (List Paragraph, 10pt)
    • Certification 2

    Pagination: Entire section stays together (certifications are typically short)
    """
    if not certifications:
        return

    # Keep header with content (entire section should stay together)
    add_section_header(doc, section_name, keep_with_next=True)

    for cert in certifications:
        add_bullet_paragraph(doc, cert, font_size=10, keep_together=True)


# ============================================================================
# CUSTOM SECTIONS BUILDERS
# ============================================================================

def add_custom_text_section(doc: Document, section_data: Dict[str, Any], section_name: str):
    """
    Add custom text section

    Args:
        doc: Document object
        section_data: Custom section data with 'content' field
        section_name: Display name for the section
    """
    if not section_data or not section_data.get('content'):
        return

    # Add section header
    add_section_header(doc, section_name, keep_with_next=True)

    # Add content paragraph
    content = section_data.get('content', '')
    para = doc.add_paragraph(sanitize_text(content))
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.left_indent = Inches(0)
    para.paragraph_format.first_line_indent = Inches(0)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    para.paragraph_format.keep_together = True
    if para.runs:
        para.runs[0].font.size = Pt(10)
        para.runs[0].font.name = 'Calibri'


def add_custom_list_section(doc: Document, section_data: Dict[str, Any], section_name: str):
    """
    Add custom list section with items containing title, subtitle, description, and bullets

    Args:
        doc: Document object
        section_data: Custom section data with 'items' field
        section_name: Display name for the section
    """
    items = section_data.get('items', [])
    if not items:
        return

    # Add section header
    add_section_header(doc, section_name, keep_with_next=True)

    for i, item in enumerate(items):
        title = item.get('title', '')
        subtitle = item.get('subtitle', '')
        description = item.get('description', '')
        bullets = item.get('bullets', [])

        # Add title and subtitle row (similar to experience)
        if title or subtitle:
            row_para = doc.add_paragraph()
            row_para.paragraph_format.space_before = Pt(0)
            row_para.paragraph_format.space_after = Pt(0)
            row_para.paragraph_format.keep_with_next = True
            row_para.paragraph_format.keep_together = True

            # Title (bold, left-aligned)
            if title:
                title_run = row_para.add_run(sanitize_text(title))
                title_run.font.size = Pt(10)
                title_run.font.name = 'Calibri'
                title_run.font.bold = True

            # Subtitle (italic, right-aligned via tab)
            if subtitle:
                if title:
                    row_para.add_run('\t')  # Tab to push subtitle to right
                subtitle_run = row_para.add_run(sanitize_text(subtitle))
                subtitle_run.font.size = Pt(10)
                subtitle_run.font.name = 'Calibri'
                subtitle_run.font.italic = True

            # Set tab stop for right alignment
            if title and subtitle:
                from docx.shared import Inches
                from docx.enum.text import WD_TAB_ALIGNMENT
                tab_stops = row_para.paragraph_format.tab_stops
                tab_stops.add_tab_stop(Inches(6.0), WD_TAB_ALIGNMENT.RIGHT)

        # Add description if present
        if description:
            desc_para = doc.add_paragraph(sanitize_text(description))
            desc_para.paragraph_format.space_before = Pt(0)
            desc_para.paragraph_format.space_after = Pt(2)
            desc_para.paragraph_format.keep_together = True
            if desc_para.runs:
                desc_para.runs[0].font.size = Pt(10)
                desc_para.runs[0].font.name = 'Calibri'

        # Add bullets if present
        if bullets:
            for bullet in bullets:
                if bullet and bullet.strip():
                    add_bullet_paragraph(doc, bullet, font_size=10, keep_together=True)

        # Add spacing between items (except after last item)
        if i < len(items) - 1:
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(4)
            spacer.paragraph_format.space_after = Pt(0)


def add_custom_simple_list_section(doc: Document, section_data: Dict[str, Any], section_name: str):
    """
    Add custom simple list section (just bullet points)

    Args:
        doc: Document object
        section_data: Custom section data with 'items' field (array of strings)
        section_name: Display name for the section
    """
    items = section_data.get('items', [])
    if not items:
        return

    # Add section header
    add_section_header(doc, section_name, keep_with_next=True)

    # Add each item as a bullet point
    for item in items:
        if item and item.strip():
            add_bullet_paragraph(doc, item, font_size=10, keep_together=True)


# ============================================================================
# MAIN GENERATION FUNCTION
# ============================================================================

def generate_resume_from_json(
    resume_json: Dict[str, Any],
    base_resume_docx: bytes = None,
    section_order: List[str] = None
) -> bytes:
    """
    Generate resume DOCX from JSON data programmatically

    Args:
        resume_json: Extracted resume data (matches ResumeData schema)
        base_resume_docx: Original resume DOCX as bytes (for style reference)
                         If None, creates new document with default styles
        section_order: Custom order for sections
                      Default: ['professional_summary', 'experience', 'projects',
                               'education', 'skills', 'certifications']

    Returns:
        bytes: Generated DOCX file as bytes
    """
    # Extract custom section names from resume_json
    custom_section_names = resume_json.get('section_names', {})

    # Default section names
    default_section_names = {
        'personal_info': 'PERSONAL INFORMATION',
        'professional_summary': 'PROFESSIONAL SUMMARY',
        'experience': 'EXPERIENCE',
        'projects': 'PROJECTS',
        'education': 'EDUCATION',
        'skills': 'TECHNICAL SKILLS',
        'certifications': 'CERTIFICATIONS',
    }

    # Helper function to get section name (custom or default)
    def get_section_name(key: str) -> str:
        """Get custom section name if exists, otherwise return default"""
        return custom_section_names.get(key, default_section_names.get(key, key.upper()))

    try:
        # Load base resume as style reference (or create new document)
        if base_resume_docx:
            logger.info("Loading style reference from base resume bytes")

            # Load the original document
            temp_doc = Document(BytesIO(base_resume_docx))

            # Clear all existing content
            for para in temp_doc.paragraphs[:]:
                p = para._element
                p.getparent().remove(p)

            # Save to bytes to refresh the document structure
            temp_output = BytesIO()
            temp_doc.save(temp_output)
            temp_output.seek(0)

            # Reload the cleared document (this preserves styles properly)
            doc = Document(temp_output)

            # Set ALL margins to 0.5 inch and page size to US Letter
            section = doc.sections[0]
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

            # Set page size to US Letter (8.5" × 11")
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)

            # Log the margins for debugging
            logger.info(f"Document margins - Top: {section.top_margin.inches}\", Bottom: {section.bottom_margin.inches}\", Left: {section.left_margin.inches}\", Right: {section.right_margin.inches}\"")

            logger.info("Cleared existing content, preserving styles and page setup")
        else:
            logger.warning("No base resume provided - creating new document with default styles")
            logger.warning("Generated document may not have exact formatting")
            doc = Document()

            # Ensure required styles exist
            # Note: This is a fallback - results won't be as good as using original resume
            pass

        # Default section order
        if section_order is None:
            section_order = [
                'professional_summary',
                'experience',
                'projects',
                'education',
                'skills',
                'certifications'
            ]

        # Build resume programmatically
        logger.info("Building resume from JSON data...")

        # 1. Header (always first)
        personal_info = resume_json.get('personal_info', {})
        add_header_section(doc, personal_info)

        # 2. Add sections in custom order
        section_builders = {
            'professional_summary': lambda: add_professional_summary(
                doc,
                resume_json.get('professional_summary', ''),
                get_section_name('professional_summary')
            ),
            'education': lambda: add_education_section(
                doc,
                resume_json.get('education', []),
                get_section_name('education')
            ),
            'experience': lambda: add_experience_section(
                doc,
                resume_json.get('experience', []),
                get_section_name('experience')
            ),
            'projects': lambda: add_projects_section(
                doc,
                resume_json.get('projects', []),
                get_section_name('projects')
            ),
            'skills': lambda: add_skills_section(
                doc,
                resume_json.get('skills', []),
                get_section_name('skills')
            ),
            'certifications': lambda: add_certifications_section(
                doc,
                resume_json.get('certifications', []),
                get_section_name('certifications')
            ),
        }

        # Build sections in order
        for section_name in section_order:
            # Check if it's a standard section
            if section_name in section_builders:
                section_builders[section_name]()
            # Check if it's a custom section
            elif section_name.startswith('custom_'):
                custom_sections = resume_json.get('custom_sections', [])
                custom_section = next((s for s in custom_sections if s.get('id') == section_name), None)

                if custom_section:
                    section_type = custom_section.get('type')
                    section_display_name = get_section_name(section_name)

                    # Call appropriate builder based on section type
                    if section_type == 'text':
                        add_custom_text_section(doc, custom_section, section_display_name)
                    elif section_type == 'list':
                        add_custom_list_section(doc, custom_section, section_display_name)
                    elif section_type == 'simple_list':
                        add_custom_simple_list_section(doc, custom_section, section_display_name)

        # Save to bytes
        output = BytesIO()
        doc.save(output)
        output.seek(0)

        logger.info("Resume generated successfully")
        return output.read()

    except Exception as e:
        logger.error(f"Resume generation failed: {e}")
        raise Exception(f"Failed to generate resume: {str(e)}")


# ============================================================================
# HELPER FUNCTION - Get Default Section Order
# ============================================================================

def get_default_section_order() -> List[str]:
    """
    Get default section order

    NOTE: personal_info is NOT included here because it's ALWAYS rendered first
    as the header section. This list only contains sections that appear after the header.
    """
    return [
        'professional_summary',
        'experience',
        'projects',
        'education',
        'skills',
        'certifications'
    ]


# ============================================================================
# MAIN - For Testing
# ============================================================================

if __name__ == '__main__':
    """Test resume generation"""

    # Sample test data
    test_data = {
        "personal_info": {
            "name": "JOHN DOE",
            "email": "john.doe@email.com",
            "phone": "(555) 123-4567",
            "location": "San Francisco, CA",
            "linkedin": "LinkedIn",
            "github": "GitHub",
            "portfolio": "Portfolio"
        },
        "professional_summary": "Experienced Software Engineer with 5+ years of experience in full-stack development.",
        "education": [
            {
                "institution": "University of Technology",
                "location": "San Francisco, CA",
                "degree": "B.S. in Computer Science",
                "gpa": "3.8",
                "graduation_date": "May 2020"
            }
        ],
        "experience": [
            {
                "company": "Tech Company Inc.",
                "title": "Software Engineer",
                "location": "San Francisco, CA",
                "start_date": "Jan 2020",
                "end_date": "Present",
                "bullets": [
                    "Developed and maintained web applications using React and Node.js",
                    "Collaborated with cross-functional teams to deliver features"
                ]
            }
        ],
        "projects": [],
        "skills": [
            {
                "category": "Languages",
                "skills": ["JavaScript", "Python", "TypeScript", "SQL"]
            }
        ],
        "certifications": []
    }

    # Generate resume with programmatic template
    output_bytes = generate_resume_from_json(test_data)

    print(f"✅ Test successful! Generated {len(output_bytes)} bytes")
    print(f"   Using programmatic template (no base resume needed)")


def generate_cover_letter_docx(cover_letter_text: str, resume_json: dict = None) -> bytes:
    """
    Generate a professional cover letter DOCX.
    Generates personal info header from resume_json, then adds LLM-generated body.

    Args:
        cover_letter_text: LLM-generated cover letter body (starting from date)
        resume_json: Resume JSON containing personal info with links

    Returns:
        bytes: DOCX file as bytes
    """
    doc = Document()

    # Set default font to Calibri with tight spacing
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Set paragraph spacing to 0 (tight spacing)
    paragraph_format = style.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing = 1.0  # Single line spacing

    # === STEP 1: Generate Personal Info Header from resume_json ===
    if resume_json:
        personal_info = resume_json.get('personal_info', {})

        # Name (left-aligned, regular weight, same size as everything else)
        name = personal_info.get('name', '')
        if name:
            name_para = doc.add_paragraph()
            name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            name_run = name_para.add_run(name)
            name_run.font.name = 'Calibri'
            name_run.font.size = Pt(11)
            name_para.paragraph_format.space_after = Pt(0)

        # Location on separate line
        if personal_info.get('location'):
            loc_para = doc.add_paragraph()
            loc_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            loc_run = loc_para.add_run(personal_info['location'])
            loc_run.font.name = 'Calibri'
            loc_run.font.size = Pt(11)
            loc_para.paragraph_format.space_after = Pt(0)

        # Email on separate line
        if personal_info.get('email'):
            email_para = doc.add_paragraph()
            email_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            email_run = email_para.add_run(personal_info['email'])
            email_run.font.name = 'Calibri'
            email_run.font.size = Pt(11)
            email_para.paragraph_format.space_after = Pt(0)

        # Phone on separate line
        if personal_info.get('phone'):
            phone_para = doc.add_paragraph()
            phone_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            phone_run = phone_para.add_run(personal_info['phone'])
            phone_run.font.name = 'Calibri'
            phone_run.font.size = Pt(11)
            phone_para.paragraph_format.space_after = Pt(0)

        # Add header links (2 per row)
        header_links = personal_info.get('header_links', [])
        if header_links:
            # Process links 2 at a time
            for i in range(0, len(header_links), 2):
                links_in_row = header_links[i:i+2]

                link_para = doc.add_paragraph()
                link_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                link_para.paragraph_format.space_after = Pt(0)

                for idx, link in enumerate(links_in_row):
                    text = link.get('text', '')
                    url = link.get('url', '')

                    if not text:
                        continue

                    # Add separator between links in the same row
                    if idx > 0:
                        sep_run = link_para.add_run(' | ')
                        sep_run.font.name = 'Calibri'
                        sep_run.font.size = Pt(11)

                    # Add hyperlink
                    if url:
                        add_professional_hyperlink(link_para, url, text, 'Calibri', 11, False)
                    else:
                        # No URL, just plain text
                        text_run = link_para.add_run(text)
                        text_run.font.name = 'Calibri'
                        text_run.font.size = Pt(11)

            # Add extra spacing after links section (before date)
            if header_links:
                link_para.paragraph_format.space_after = Pt(12)

    # === STEP 2: Add LLM-generated body (starting from date) ===
    # The cover letter text only contains: date, company info, letter body
    # Personal info header was already generated above from resume_json
    paragraphs_text = cover_letter_text.split('\n')

    for para_text in paragraphs_text:
        if not para_text.strip():
            # Empty line - add minimal spacing (6pt instead of full paragraph)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            continue

        # Add paragraph as-is (LLM-generated body content)
        p = doc.add_paragraph(para_text)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Apply Calibri font and tight spacing
        for run in p.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(11)

        # Set tight line spacing for this paragraph
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    # Save to BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


def add_professional_hyperlink(paragraph, url, text, font_name='Calibri', font_size=11, underline=False):
    """
    Add a hyperlink to a paragraph for cover letters.
    The text will be displayed in the same font and color as normal text (black).

    Args:
        paragraph: The paragraph to add hyperlink to
        url: The URL to link to
        text: The display text for the link
        font_name: Font name (default: Calibri)
        font_size: Font size in points (default: 11)
        underline: Whether to underline (default: False for professional look)
    """
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a new run object (a wrapper over a 'w:r' element)
    new_run = OxmlElement('w:r')

    # Set font properties
    rPr = OxmlElement('w:rPr')

    # Font name
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rPr.append(rFonts)

    # Font size
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(font_size * 2))  # Font size is in half-points
    rPr.append(sz)

    # Color - use black (same as normal text)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '000000')  # Black
    rPr.append(color)

    # Remove underline for professional look (optional)
    if not underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'none')
        rPr.append(u)

    new_run.append(rPr)

    # Join all the xml elements together
    new_run.text = text
    hyperlink.append(new_run)

    # Append the hyperlink to the paragraph
    paragraph._p.append(hyperlink)

    return hyperlink
