"""
Resume Extraction Service - Hybrid Multi-Format Support
Supports DOCX, PDF, and Image files with intelligent OCR fallback
"""

from openai import OpenAI
from pydantic import BaseModel, Field
from typing import List, Optional, Dict, Callable
from docx import Document
from io import BytesIO
import logging
from config.settings import settings
import fitz  # PyMuPDF
import docx2txt
from pdf2image import convert_from_bytes
import pytesseract
from PIL import Image
import json

# Configure logging
logger = logging.getLogger(__name__)


# ============================================================================
# SCHEMA DEFINITIONS (Pydantic Models)
# ============================================================================

class HeaderLink(BaseModel):
    """Generic header link - can be any type (LinkedIn, GitHub, etc.)"""
    text: str  # Display text (e.g., "LinkedIn", "Portfolio")
    url: Optional[str] = None  # Actual URL if available


class PersonalInfo(BaseModel):
    """Personal information"""
    name: str
    current_role: Optional[str] = None  # Job title/role if mentioned in header
    email: Optional[str] = None
    phone: Optional[str] = None
    location: Optional[str] = None
    header_links: List[HeaderLink] = []  # Generic dynamic links


class ExperienceEntry(BaseModel):
    """Work experience entry"""
    company: str
    title: str
    location: Optional[str] = None
    start_date: str
    end_date: str  # "Present" if current
    bullets: List[str]


class EducationEntry(BaseModel):
    """Education entry"""
    degree: str
    institution: str
    location: Optional[str] = None
    graduation_date: str
    gpa: Optional[str] = None
    gpa_out_of: Optional[str] = None  # e.g., "4" for "3.5/4" format


class SkillCategory(BaseModel):
    """Skills grouped by category"""
    category: str  # e.g., "Languages", "Frontend", "Tools"
    skills: List[str]


class ProjectEntry(BaseModel):
    """Project entry"""
    name: str
    bullets: List[str] = []  # Changed from description to bullets for consistency with experience
    description: Optional[str] = None  # Keep for backward compatibility
    technologies: List[str] = []
    link: Optional[str] = None
    start_date: Optional[str] = None  # Added for consistency
    end_date: Optional[str] = None  # Added for consistency


class ResumeData(BaseModel):
    """Complete resume structure - OpenAI will fill this"""
    personal_info: PersonalInfo
    professional_summary: Optional[str] = None
    experience: List[ExperienceEntry] = []
    education: List[EducationEntry] = []
    skills: List[SkillCategory] = []
    projects: List[ProjectEntry] = []
    certifications: List[str] = []


# ============================================================================
# EXTRACTOR SERVICE WITH HYBRID MULTI-FORMAT SUPPORT
# ============================================================================

class ResumeExtractor:
    """
    Hybrid Resume Extractor - Supports DOCX, PDF, and Images

    Features:
    - Fast text extraction for standard resumes
    - OCR fallback for image-based resumes
    - Status callbacks for user feedback
    - Handles mixed content (text + images)

    Usage:
        extractor = ResumeExtractor(api_key)
        resume_json = extractor.extract(
            file_bytes,
            file_type='docx',
            status_callback=lambda msg: print(msg)
        )
    """

    # Supported file formats
    SUPPORTED_FORMATS = {
        'docx': ['.docx', '.doc'],
        'pdf': ['.pdf'],
        'image': ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif']
    }

    # Minimum characters to consider extraction successful
    MIN_EXTRACTION_THRESHOLD = 100

    def __init__(self, api_key: str = None):
        """Initialize with OpenAI API key"""
        self.api_key = api_key or settings.OPENAI_API_KEY

        if not self.api_key:
            raise ValueError(
                "OpenAI API key required. Set OPENAI_API_KEY in .env file"
            )

        self.client = OpenAI(api_key=self.api_key)

    def validate_file_format(self, filename: str) -> str:
        """
        Validate file format and return file type category

        Returns: 'docx', 'pdf', or 'image'
        Raises: ValueError if unsupported format
        """
        filename_lower = filename.lower()

        for file_type, extensions in self.SUPPORTED_FORMATS.items():
            if any(filename_lower.endswith(ext) for ext in extensions):
                return file_type

        # If not supported, raise error
        supported_list = []
        for exts in self.SUPPORTED_FORMATS.values():
            supported_list.extend(exts)

        raise ValueError(
            f"Unsupported file format. Please upload one of: {', '.join(supported_list)}"
        )

    def _extract_text_from_docx(self, docx_bytes: bytes) -> str:
        """Extract text from DOCX using multiple methods, including hyperlinks"""
        try:
            text_parts = []
            hyperlinks = []

            # Method 1: python-docx (preserves structure)
            doc = Document(BytesIO(docx_bytes))

            # Extract from text boxes first (often contains header)
            from docx.text.paragraph import Paragraph
            for element in doc.element.body:
                for shape in element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}txbxContent'):
                    for p_elem in shape.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'):
                        para = Paragraph(p_elem, doc)
                        text = para.text.strip()
                        if text:
                            text_parts.append(text)

            # Extract paragraphs and hyperlinks
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    text_parts.append(text)

                # Extract hyperlinks from this paragraph
                for run in para.runs:
                    if hasattr(run, '_element'):
                        # Look for hyperlink elements in the run
                        for elem in run._element.iterchildren():
                            if elem.tag.endswith('hyperlink'):
                                # Get the relationship ID
                                r_id = elem.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                                if r_id:
                                    try:
                                        url = doc.part.rels[r_id].target_ref
                                        link_text = ''.join(node.text for node in elem.iter() if hasattr(node, 'text') and node.text)
                                        if url and link_text:
                                            hyperlinks.append(f"{link_text.strip()}: {url}")
                                        elif url:
                                            hyperlinks.append(f"Link: {url}")
                                    except:
                                        pass

            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            text_parts.append(cell_text)

            extracted = '\n'.join(text_parts)

            # Add hyperlinks section if any found
            if hyperlinks:
                extracted += "\n\n--- Hyperlinks ---\n" + "\n".join(hyperlinks)

            # Method 2: If python-docx failed, try docx2txt (sometimes catches more)
            if len(extracted.strip()) < self.MIN_EXTRACTION_THRESHOLD:
                logger.info("python-docx extraction insufficient, trying docx2txt...")
                extracted = docx2txt.process(BytesIO(docx_bytes))

            return extracted

        except Exception as e:
            logger.error(f"DOCX text extraction failed: {e}")
            return ""

    def _extract_text_from_pdf(self, pdf_bytes: bytes) -> str:
        """Extract text from PDF using PyMuPDF, including hyperlinks"""
        try:
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            text_parts = []

            for page in doc:
                # Extract visible text
                text = page.get_text()
                if text:
                    text_parts.append(text)

                # Extract hyperlinks from the page
                links = page.get_links()
                if links:
                    link_texts = []
                    for link in links:
                        # Check if link has a URI (external link)
                        if 'uri' in link and link['uri']:
                            uri = link['uri']
                            # Try to get the link text by extracting text from the link rectangle
                            rect = link.get('from', None)
                            if rect:
                                try:
                                    link_text = page.get_text("text", clip=rect).strip()
                                    if link_text:
                                        link_texts.append(f"{link_text}: {uri}")
                                    else:
                                        link_texts.append(f"Link: {uri}")
                                except:
                                    link_texts.append(f"Link: {uri}")
                            else:
                                link_texts.append(f"Link: {uri}")

                    # Add links section to extracted text
                    if link_texts:
                        text_parts.append("\n--- Hyperlinks ---\n" + "\n".join(link_texts))

            doc.close()
            return '\n'.join(text_parts)

        except Exception as e:
            logger.error(f"PDF text extraction failed: {e}")
            return ""

    def _extract_text_with_ocr(self, file_bytes: bytes, file_type: str) -> str:
        """
        Extract text using OCR (works for images and image-based PDFs/DOCX)

        Args:
            file_bytes: File content as bytes
            file_type: 'docx', 'pdf', or 'image'

        Returns:
            Extracted text string
        """
        try:
            images = []

            if file_type == 'image':
                # Direct image file
                image = Image.open(BytesIO(file_bytes))
                images = [image]

            elif file_type == 'pdf':
                # Convert PDF pages to images
                logger.info("Converting PDF to images for OCR...")
                images = convert_from_bytes(file_bytes, dpi=300)

            elif file_type == 'docx':
                # Convert DOCX to PDF first, then to images
                logger.info("Converting DOCX to PDF for OCR...")
                import subprocess
                import tempfile
                import os

                # Save DOCX to temp file
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_docx:
                    tmp_docx.write(file_bytes)
                    docx_path = tmp_docx.name

                # Convert to PDF using LibreOffice
                output_dir = tempfile.gettempdir()
                try:
                    subprocess.run(
                        ['soffice', '--headless', '--convert-to', 'pdf', '--outdir', output_dir, docx_path],
                        capture_output=True,
                        timeout=30,
                        check=True
                    )

                    # Get PDF path
                    pdf_path = docx_path.replace('.docx', '.pdf')

                    # Read PDF and convert to images
                    with open(pdf_path, 'rb') as pdf_file:
                        pdf_bytes_converted = pdf_file.read()
                        images = convert_from_bytes(pdf_bytes_converted, dpi=300)

                    # Clean up
                    os.remove(pdf_path)

                except Exception as e:
                    logger.error(f"DOCX to PDF conversion failed: {e}")
                    raise
                finally:
                    # Clean up temp DOCX
                    os.remove(docx_path)

            # Perform OCR on all images
            text_parts = []
            for i, image in enumerate(images):
                logger.info(f"Performing OCR on page/image {i+1}/{len(images)}...")
                text = pytesseract.image_to_string(image, config='--psm 1')
                text_parts.append(text)

            extracted_text = '\n\n'.join(text_parts)
            logger.info(f"OCR extracted {len(extracted_text)} characters")

            return extracted_text

        except Exception as e:
            logger.error(f"OCR extraction failed: {e}")
            raise Exception(f"OCR extraction failed: {str(e)}")

    def _validate_is_resume(self, text: str) -> tuple[bool, str]:
        """
        Validate if extracted text is actually a resume/CV

        Uses GPT-4o-mini for fast, cheap validation before expensive structured parsing.

        Args:
            text: Extracted text from document

        Returns:
            (is_valid, reason) - Boolean and explanation string
        """
        logger.info("Validating document is a resume...")

        # Only analyze first 3000 characters for speed and cost
        text_sample = text[:3000]

        prompt = f"""You are a document classifier. Determine if the following text is a RESUME/CV.

A valid resume typically contains:
- Personal/contact information (name, email, phone, or location)
- Work experience OR education
- Skills OR qualifications

INVALID documents (reject these):
- Cover letters (focus on job application, not qualifications)
- Job descriptions or job postings
- Random documents, essays, articles, books
- Marketing materials, brochures, advertisements
- Blank or gibberish content
- Forms, templates, invoices, or receipts
- Code files, logs, or technical documentation

Analyze this text and respond in JSON format:
{{
    "is_resume": true/false,
    "reason": "brief explanation (1-2 sentences)",
    "confidence": "high/medium/low"
}}

TEXT TO ANALYZE:
{text_sample}
"""

        try:
            response = self.client.chat.completions.create(
                model="gpt-4o-mini",  # Cheaper + faster for validation (~$0.0001 per call)
                messages=[
                    {
                        "role": "system",
                        "content": "You are a document classifier that validates if a document is a resume/CV. Be strict but fair."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                response_format={"type": "json_object"},
                temperature=0.1,
                max_tokens=200
            )

            result = json.loads(response.choices[0].message.content)
            is_resume = result.get("is_resume", False)
            reason = result.get("reason", "Unknown reason")
            confidence = result.get("confidence", "low")

            logger.info(f"Validation result: is_resume={is_resume}, confidence={confidence}, reason={reason}")

            return is_resume, reason

        except Exception as e:
            logger.error(f"Resume validation failed: {e}")
            # On error, fail open (allow upload to proceed)
            # Better to allow a few false positives than block legitimate resumes
            return True, "Validation check encountered an error, allowing upload"

    def _clean_extracted_data(self, data: dict) -> dict:
        """
        Post-process extracted data to remove duplicates and clean up

        Args:
            data: Raw extracted resume data

        Returns:
            Cleaned resume data
        """
        # Clean header_links: remove email/phone duplicates
        if 'personal_info' in data and 'header_links' in data['personal_info']:
            email = data['personal_info'].get('email', '') or ''
            phone = data['personal_info'].get('phone', '') or ''

            # Ensure email and phone are strings
            email = str(email).lower() if email else ''
            phone = str(phone) if phone else ''

            cleaned_links = []
            for link in data['personal_info']['header_links']:
                link_text = str(link.get('text', '') or '').lower()
                link_url = str(link.get('url', '') or '').lower()

                # Skip if it's email or phone
                if email and (email in link_text or email in link_url):
                    continue
                if phone and (phone in link_text or phone in link_url):
                    continue

                cleaned_links.append(link)

            data['personal_info']['header_links'] = cleaned_links

        # Deduplicate skills within each category
        if 'skills' in data:
            for skill_category in data['skills']:
                if 'skills' in skill_category and isinstance(skill_category['skills'], list):
                    # Remove duplicates while preserving order
                    seen = set()
                    unique_skills = []
                    for skill in skill_category['skills']:
                        skill_lower = skill.lower().strip()
                        if skill_lower not in seen:
                            seen.add(skill_lower)
                            unique_skills.append(skill.strip())
                    skill_category['skills'] = unique_skills

        return data

    def _parse_with_llm(self, text: str) -> dict:
        """Parse extracted text with OpenAI LLM"""
        logger.info("Sending to OpenAI for structured parsing...")

        prompt = f"""
Extract all information from this resume into structured format.

INSTRUCTIONS:
- Identify sections by context (not just keywords)
- Extract ALL content exactly as written
- Maintain chronological order (newest first)
- For skills, group by logical categories (remove duplicates within each category)
- Preserve dates exactly (e.g., "Sept 2025", "Present")
- For personal_info:
  * Extract current_role/job title if mentioned in the header
  * Clean up the job title to be professional and standard:
    - Remove marketing buzzwords like "AI-driven", "Innovative", "Dynamic", "Passionate"
    - Extract the core role (e.g., "Full Stack Developer", "Software Engineer", "Data Analyst")
    - Acceptable prefixes: "Senior", "Junior", "Lead", "Principal", "Staff"
    - Examples: "AI-driven Full Stack Developer" → "Full Stack Developer"
                "Innovative Software Engineer" → "Software Engineer"
                "Passionate Data Scientist" → "Data Scientist"
  * DO NOT include email/phone in header_links (they have dedicated fields)
  * For header_links: Extract social/professional links (LinkedIn, GitHub, Portfolio, etc.)
  * If you see a "--- Hyperlinks ---" section, use it to extract both link text AND URL
  * Format: text field should be the link display text, url field should be the actual URL
  * Keep link text clean and simple (e.g., "LinkedIn" not "LinkedinLinkedinLinkedin")
  * Avoid repeating text multiple times in link names
- For projects:
  * Extract each bullet point as a separate item in the bullets array
  * DO NOT concatenate bullets into a single description string
  * Each bullet should be a distinct array element
  * Include start_date and end_date if mentioned

RESUME:
{text}
"""

        completion = self.client.beta.chat.completions.parse(
            model="gpt-4o-2024-08-06",
            messages=[
                {
                    "role": "system",
                    "content": "You are an expert resume parser. Extract information accurately into structured JSON."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            response_format=ResumeData,
            temperature=0.1
        )

        resume_data = completion.choices[0].message.parsed
        result = resume_data.model_dump()

        # Clean up the data
        result = self._clean_extracted_data(result)

        logger.info("Resume parsed successfully by LLM")
        return result

    def extract(
        self,
        file_bytes: bytes,
        filename: str = "resume.docx",
        status_callback: Optional[Callable[[str], None]] = None
    ) -> dict:
        """
        Main extraction method with hybrid approach

        Args:
            file_bytes: File content as bytes
            filename: Original filename (used to detect format)
            status_callback: Optional callback for status updates

        Returns:
            dict: Structured resume data matching ResumeData schema

        Raises:
            ValueError: If file format is unsupported
            Exception: If extraction fails completely
        """

        def send_status(message: str):
            """Send status update via callback"""
            logger.info(message)
            if status_callback:
                status_callback(message)

        try:
            # Step 1: Validate file format
            send_status("Validating file format...")
            file_type = self.validate_file_format(filename)
            send_status(f"File type detected: {file_type.upper()}")

            # Step 2: Try fast text extraction first
            extracted_text = ""

            if file_type == 'docx':
                send_status("Extracting text from DOCX...")
                extracted_text = self._extract_text_from_docx(file_bytes)

            elif file_type == 'pdf':
                send_status("Extracting text from PDF...")
                extracted_text = self._extract_text_from_pdf(file_bytes)

            elif file_type == 'image':
                # For images, skip straight to OCR
                send_status("Image file detected, using OCR...")
                extracted_text = ""  # Force OCR

            # Step 3: Check if extraction was successful
            if len(extracted_text.strip()) < self.MIN_EXTRACTION_THRESHOLD:
                send_status(
                    "Switching to OCR for better extraction... This may take 5-10 seconds."
                )
                extracted_text = self._extract_text_with_ocr(file_bytes, file_type)

                # Validate OCR extraction
                if len(extracted_text.strip()) < self.MIN_EXTRACTION_THRESHOLD:
                    raise Exception(
                        "Unable to extract meaningful text from the document. "
                        "Please ensure the file contains readable text or try a different format."
                    )

            send_status(f"Successfully extracted {len(extracted_text)} characters")

            # Step 3.5: Validate document is actually a resume
            send_status("Validating document is a resume...")
            is_valid, reason = self._validate_is_resume(extracted_text)

            if not is_valid:
                error_msg = f"Document validation failed: {reason}"
                logger.warning(error_msg)
                raise ValueError("Please upload a Valid resume")

            send_status("Document validated as resume ✓")

            # Step 4: Parse with LLM
            send_status("Analyzing resume content with AI...")
            result = self._parse_with_llm(extracted_text)

            send_status("Resume extraction completed successfully!")
            return result

        except ValueError as e:
            # File format validation error
            logger.error(f"File format error: {e}")
            raise

        except Exception as e:
            logger.error(f"Resume extraction failed: {e}")
            raise Exception(f"Failed to extract resume: {str(e)}")


# ============================================================================
# HELPER FUNCTION (Simple API)
# ============================================================================

def extract_resume(
    file_bytes: bytes,
    filename: str = "resume.docx",
    api_key: str = None,
    status_callback: Optional[Callable[[str], None]] = None
) -> dict:
    """
    Simple function to extract resume with hybrid approach

    Usage:
        def on_status(message):
            print(f"Status: {message}")

        resume_json = extract_resume(
            file_bytes,
            filename="resume.pdf",
            status_callback=on_status
        )
    """
    extractor = ResumeExtractor(api_key)
    return extractor.extract(file_bytes, filename, status_callback)
