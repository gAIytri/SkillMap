"""
DOCX Recreation Service
Recreates DOCX from original file + JSON modifications with exact styling preservation
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from io import BytesIO
from typing import Dict, Any, List, Optional
import logging
import re
import json

logger = logging.getLogger(__name__)


def recreate_docx_from_json(original_docx_bytes: bytes, resume_json: Dict[str, Any]) -> bytes:
    """
    Recreate DOCX from original file with updates from JSON
    Preserves all styling: fonts, colors, spacing, formatting

    Args:
        original_docx_bytes: Original DOCX file bytes
        resume_json: Extracted/modified JSON data

    Returns:
        bytes: Recreated DOCX file with updated content
    """
    try:
        logger.info("Recreating DOCX from JSON with styling preservation...")

        # Load original document
        doc = Document(BytesIO(original_docx_bytes))

        # Apply JSON changes to document
        update_personal_info(doc, resume_json.get('personal_info', {}))
        update_professional_summary(doc, resume_json.get('professional_summary'))
        update_experience(doc, resume_json.get('experience', []))
        update_education(doc, resume_json.get('education', []))
        update_skills(doc, resume_json.get('skills', []))
        update_projects(doc, resume_json.get('projects', []))
        update_certifications(doc, resume_json.get('certifications', []))

        # Save to bytes
        output = BytesIO()
        doc.save(output)
        output.seek(0)

        logger.info("DOCX recreated successfully from JSON")
        return output.getvalue()

    except Exception as e:
        logger.error(f"DOCX recreation failed: {e}")
        raise Exception(f"Failed to recreate DOCX: {str(e)}")


def update_personal_info(doc: Document, personal_info: Dict[str, Any]):
    """Update personal information section (usually at top)"""
    if not personal_info:
        return

    # Find and update name (usually first paragraph with large font or Title style)
    for para in doc.paragraphs[:5]:  # Check first 5 paragraphs
        if para.style.name == 'Title' or (para.runs and len(para.runs) > 0 and
                                          para.runs[0].font.size and
                                          para.runs[0].font.size.pt > 14):
            # This is likely the name
            if personal_info.get('name'):
                update_paragraph_text(para, personal_info['name'])
            break

    # Find and update contact info (email, phone, location)
    # Usually in first few paragraphs, look for patterns
    for para in doc.paragraphs[:10]:
        text = para.text.lower()

        # Update email
        if '@' in text and personal_info.get('email'):
            new_text = re.sub(r'[\w\.-]+@[\w\.-]+\.\w+', personal_info['email'], para.text)
            update_paragraph_text(para, new_text)

        # Update phone
        if re.search(r'\(\d{3}\)|\d{3}[-.]?\d{3}[-.]?\d{4}', text) and personal_info.get('phone'):
            new_text = re.sub(r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', personal_info['phone'], para.text)
            update_paragraph_text(para, new_text)

        # Update LinkedIn
        if 'linkedin' in text and personal_info.get('linkedin'):
            new_text = re.sub(r'linkedin\.com/[\w\-/]+', personal_info['linkedin'], para.text, flags=re.IGNORECASE)
            update_paragraph_text(para, new_text)

        # Update GitHub
        if 'github' in text and personal_info.get('github'):
            new_text = re.sub(r'github\.com/[\w\-/]+', personal_info['github'], para.text, flags=re.IGNORECASE)
            update_paragraph_text(para, new_text)


def update_professional_summary(doc: Document, summary: Optional[str]):
    """Update professional summary section"""
    if not summary:
        return

    # Find summary section (look for keywords like "Summary", "Profile", "About")
    summary_keywords = ['summary', 'profile', 'about', 'objective']

    for i, para in enumerate(doc.paragraphs):
        text = para.text.lower().strip()

        # Check if this is a summary heading
        if any(keyword in text for keyword in summary_keywords) and len(text) < 50:
            # Next paragraph(s) likely contain the summary
            # Replace content while preserving formatting
            if i + 1 < len(doc.paragraphs):
                update_paragraph_text(doc.paragraphs[i + 1], summary)
            break


def update_experience(doc: Document, experience_list: List[Dict[str, Any]]):
    """Update experience section with bullet points"""
    if not experience_list:
        logger.warning("No experience data to update")
        return

    # Debug: Log the structure of experience data
    logger.info(f"Experience list structure: {json.dumps(experience_list[:1], indent=2)[:500]}...")

    # Find experience section
    exp_start_idx = find_section_start(doc, ['experience', 'work history', 'employment', 'professional experience'])
    if exp_start_idx == -1:
        logger.warning("Could not find experience section in document")
        return

    logger.info(f"Found experience section at index {exp_start_idx}")

    # Find where experience section ends
    exp_end_idx = find_next_section(doc, exp_start_idx)
    logger.info(f"Experience section ends at index {exp_end_idx}")

    # Get the experience paragraphs
    exp_paragraphs = doc.paragraphs[exp_start_idx + 1:exp_end_idx]
    logger.info(f"Processing {len(exp_paragraphs)} paragraphs in experience section")

    # Group paragraphs by job (company/title lines followed by bullets)
    current_exp_idx = 0
    i = 0

    while i < len(exp_paragraphs) and current_exp_idx < len(experience_list):
        para = exp_paragraphs[i]
        exp_data = experience_list[current_exp_idx]

        # Check if this is a job title/company line (not a bullet)
        if not is_bullet_point(para):
            # Update company and title line
            new_text = f"{exp_data['title']} at {exp_data['company']}"
            if exp_data.get('location'):
                new_text += f" | {exp_data['location']}"
            logger.info(f"Updating job title: {new_text}")
            update_paragraph_text(para, new_text)

            # Next line might be dates
            if i + 1 < len(exp_paragraphs):
                date_para = exp_paragraphs[i + 1]
                if not is_bullet_point(date_para):
                    new_date = f"{exp_data['start_date']} - {exp_data['end_date']}"
                    update_paragraph_text(date_para, new_date)
                    i += 1

            i += 1

            # Update bullets
            bullet_idx = 0
            bullets_raw = exp_data.get('bullets', [])

            # Handle different bullet formats
            bullets_to_update = []

            if isinstance(bullets_raw, str):
                # String format - split by sentence delimiters
                logger.info(f"Bullets is string format, splitting...")
                # Try newlines first
                if '\n' in bullets_raw:
                    bullets_to_update = [b.strip() for b in bullets_raw.split('\n') if b.strip()]
                # Otherwise split by periods
                elif '. ' in bullets_raw:
                    bullets_to_update = [b.strip() for b in bullets_raw.split('. ') if b.strip()]
                else:
                    bullets_to_update = [bullets_raw.strip()]

            elif isinstance(bullets_raw, list):
                if len(bullets_raw) == 1 and isinstance(bullets_raw[0], str) and '. ' in bullets_raw[0]:
                    # Single element array containing paragraph - split it
                    logger.info(f"Bullets is single-element array with paragraph, splitting...")
                    single_text = bullets_raw[0]
                    bullets_to_update = [b.strip() for b in single_text.split('. ') if b.strip()]
                else:
                    # Normal array of bullets
                    bullets_to_update = [b.strip() if isinstance(b, str) else str(b) for b in bullets_raw]

            logger.info(f"Processing {len(bullets_to_update)} bullets for {exp_data['company']}")
            logger.info(f"First bullet preview: {bullets_to_update[0][:80] if bullets_to_update else 'N/A'}...")

            while i < len(exp_paragraphs) and is_bullet_point(exp_paragraphs[i]) and bullet_idx < len(bullets_to_update):
                bullet_text = bullets_to_update[bullet_idx].strip()
                # Remove trailing period if exists (will be added back by formatting)
                if bullet_text.endswith('.'):
                    bullet_text = bullet_text[:-1]

                logger.info(f"Updating bullet {bullet_idx + 1}: {bullet_text[:50]}...")
                update_paragraph_text(exp_paragraphs[i], bullet_text)
                bullet_idx += 1
                i += 1

            logger.info(f"Updated {bullet_idx} bullets for {exp_data['company']}")
            current_exp_idx += 1
        else:
            i += 1

    logger.info(f"Completed updating {current_exp_idx} experience entries")


def update_education(doc: Document, education_list: List[Dict[str, Any]]):
    """Update education section"""
    if not education_list:
        return

    # Find education section
    edu_start_idx = find_section_start(doc, ['education', 'academic', 'qualifications'])
    if edu_start_idx == -1:
        return

    edu_end_idx = find_next_section(doc, edu_start_idx)
    edu_paragraphs = doc.paragraphs[edu_start_idx + 1:edu_end_idx]

    # Update education entries
    current_edu_idx = 0
    i = 0

    while i < len(edu_paragraphs) and current_edu_idx < len(education_list):
        para = edu_paragraphs[i]
        edu_data = education_list[current_edu_idx]

        if not is_bullet_point(para):
            # Update degree line
            update_paragraph_text(para, edu_data['degree'])
            i += 1

            # Update institution and date
            if i < len(edu_paragraphs):
                inst_text = edu_data['institution']
                if edu_data.get('graduation_date'):
                    inst_text += f" | {edu_data['graduation_date']}"
                if edu_data.get('gpa'):
                    inst_text += f" | GPA: {edu_data['gpa']}"
                update_paragraph_text(edu_paragraphs[i], inst_text)
                i += 1

            current_edu_idx += 1
        else:
            i += 1


def update_skills(doc: Document, skills_list: List[Dict[str, Any]]):
    """Update skills section"""
    if not skills_list:
        return

    # Find skills section
    skills_start_idx = find_section_start(doc, ['skills', 'technical skills', 'competencies', 'technologies'])
    if skills_start_idx == -1:
        return

    skills_end_idx = find_next_section(doc, skills_start_idx)
    skills_paragraphs = doc.paragraphs[skills_start_idx + 1:skills_end_idx]

    # Update skills (usually "Category: skill1, skill2, skill3")
    for i, skill_cat in enumerate(skills_list):
        if i < len(skills_paragraphs):
            skills_text = f"{skill_cat['category']}: {', '.join(skill_cat['skills'])}"
            update_paragraph_text(skills_paragraphs[i], skills_text)


def update_projects(doc: Document, projects_list: List[Dict[str, Any]]):
    """Update projects section"""
    if not projects_list:
        logger.warning("No projects data to update")
        return

    proj_start_idx = find_section_start(doc, ['projects', 'personal projects', 'portfolio'])
    if proj_start_idx == -1:
        logger.warning("Could not find projects section in document")
        return

    logger.info(f"Found projects section at index {proj_start_idx}")

    proj_end_idx = find_next_section(doc, proj_start_idx)
    proj_paragraphs = doc.paragraphs[proj_start_idx + 1:proj_end_idx]
    logger.info(f"Processing {len(proj_paragraphs)} paragraphs in projects section")

    current_proj_idx = 0
    i = 0

    while i < len(proj_paragraphs) and current_proj_idx < len(projects_list):
        para = proj_paragraphs[i]
        proj_data = projects_list[current_proj_idx]

        if not is_bullet_point(para):
            # Update project name
            logger.info(f"Updating project name: {proj_data['name']}")
            update_paragraph_text(para, proj_data['name'])
            i += 1

            # Update description (or bullets)
            if i < len(proj_paragraphs):
                # Check if next paragraphs are bullets (some resumes have bullet points for projects)
                if is_bullet_point(proj_paragraphs[i]):
                    # Project has bullet points - split description appropriately
                    description = proj_data['description']
                    project_bullets = []

                    # Handle different description formats
                    if '\n' in description:
                        # Split by newlines
                        project_bullets = [b.strip() for b in description.split('\n') if b.strip()]
                    elif '. ' in description and len(description) > 100:
                        # Long paragraph - split by periods
                        project_bullets = [b.strip() for b in description.split('. ') if b.strip()]
                    else:
                        # Single description
                        project_bullets = [description]

                    logger.info(f"Updating {len(project_bullets)} bullets for project {proj_data['name']}")
                    logger.info(f"First project bullet: {project_bullets[0][:80] if project_bullets else 'N/A'}...")

                    bullet_idx = 0
                    while i < len(proj_paragraphs) and is_bullet_point(proj_paragraphs[i]) and bullet_idx < len(project_bullets):
                        bullet_text = project_bullets[bullet_idx].strip()
                        # Remove trailing period if exists
                        if bullet_text.endswith('.'):
                            bullet_text = bullet_text[:-1]
                        update_paragraph_text(proj_paragraphs[i], bullet_text)
                        bullet_idx += 1
                        i += 1
                else:
                    # Single description paragraph
                    logger.info(f"Updating project description: {proj_data['description'][:50]}...")
                    update_paragraph_text(proj_paragraphs[i], proj_data['description'])
                    i += 1

            # Update technologies
            if proj_data.get('technologies') and i < len(proj_paragraphs):
                tech_text = f"Technologies: {', '.join(proj_data['technologies'])}"
                update_paragraph_text(proj_paragraphs[i], tech_text)
                i += 1

            current_proj_idx += 1
        else:
            i += 1

    logger.info(f"Completed updating {current_proj_idx} project entries")


def update_certifications(doc: Document, certifications_list: List[str]):
    """Update certifications section"""
    if not certifications_list:
        return

    cert_start_idx = find_section_start(doc, ['certifications', 'certificates', 'licenses'])
    if cert_start_idx == -1:
        return

    cert_end_idx = find_next_section(doc, cert_start_idx)
    cert_paragraphs = doc.paragraphs[cert_start_idx + 1:cert_end_idx]

    for i, cert in enumerate(certifications_list):
        if i < len(cert_paragraphs):
            update_paragraph_text(cert_paragraphs[i], cert)


# Helper functions

def update_paragraph_text(para, new_text: str):
    """Update paragraph text while preserving all formatting including hyperlinks"""
    if not para.runs:
        para.add_run(new_text)
        return

    # Check if paragraph contains hyperlinks
    hyperlinks = para._element.findall('.//' + qn('w:hyperlink'))

    if hyperlinks:
        # Paragraph has hyperlinks - preserve them
        update_paragraph_with_hyperlinks(para, new_text, hyperlinks)
    else:
        # No hyperlinks - simple update
        update_paragraph_simple(para, new_text)


def update_paragraph_with_hyperlinks(para, new_text: str, hyperlinks):
    """Update paragraph that contains hyperlinks - preserve the hyperlink structure"""
    # Strategy: DON'T modify hyperlinks - they're usually email/LinkedIn/GitHub links
    # Just skip updating if paragraph has hyperlinks to preserve clickability
    # The text will remain as-is with working hyperlinks

    # For now, we'll just leave hyperlink paragraphs unchanged to preserve formatting
    # This is safer than trying to update text inside hyperlink XML
    pass


def update_paragraph_simple(para, new_text: str):
    """Update paragraph without hyperlinks - simple text replacement"""
    # Keep first run's formatting
    first_run = para.runs[0]

    # Store formatting
    font_name = first_run.font.name
    font_size = first_run.font.size
    bold = first_run.bold
    italic = first_run.italic
    underline = first_run.underline
    color = first_run.font.color.rgb if first_run.font.color.rgb else None

    # Check if this is a bullet point - preserve bullet character
    is_bullet = is_bullet_point(para)
    bullet_char = ''
    if is_bullet:
        # Extract bullet character (•, -, *, etc.)
        text = para.text.strip()
        if text and text[0] in ['•', '-', '*', '○', '▪']:
            bullet_char = text[0] + ' '

    # Remove all runs
    for run in list(para.runs):
        run._element.getparent().remove(run._element)

    # Add new run with preserved formatting
    # Add bullet character back if it was a bullet point
    final_text = bullet_char + new_text if is_bullet and bullet_char else new_text
    new_run = para.add_run(final_text)
    new_run.font.name = font_name
    if font_size:
        new_run.font.size = font_size
    new_run.bold = bold
    new_run.italic = italic
    new_run.underline = underline
    if color:
        new_run.font.color.rgb = color


def find_section_start(doc: Document, keywords: List[str]) -> int:
    """Find the index of paragraph that starts a section - improved matching"""
    for i, para in enumerate(doc.paragraphs):
        text = para.text.lower().strip()

        # More flexible matching - check if ANY keyword appears in text
        has_keyword = any(keyword.lower() in text for keyword in keywords)

        if has_keyword and len(text) < 100:  # Increased from 50 to handle longer headings
            # Multiple ways a heading can be identified
            is_heading = (
                para.style.name.startswith('Heading') or
                para.style.name == 'Title' or
                (para.runs and len(para.runs) > 0 and para.runs[0].bold) or
                (para.runs and len(para.runs) > 0 and para.runs[0].font.size and para.runs[0].font.size.pt > 12) or
                text.isupper()  # All caps headings
            )

            if is_heading:
                logger.info(f"Found section '{text}' at index {i} (style: {para.style.name})")
                return i

    logger.warning(f"Could not find section with keywords: {keywords}")
    return -1


def find_next_section(doc: Document, start_idx: int) -> int:
    """Find where the next section starts (or end of document)"""
    for i in range(start_idx + 1, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        # Check if this is a section heading
        if (para.style.name.startswith('Heading') or
            (para.runs and para.runs[0].bold and len(para.text.strip()) < 50) or
            para.style.name == 'Title'):
            return i
    return len(doc.paragraphs)


def is_bullet_point(para) -> bool:
    """Check if paragraph is a bullet point - improved detection"""
    if not para.text or not para.text.strip():
        return False

    text_start = para.text.strip()[:3]

    # Check for list style
    if para.style.name.startswith('List'):
        return True

    # Check for bullet characters at start
    bullet_chars = ['•', '-', '*', '○', '▪', '·', '◦', '▫', '►', '➢', '⬩']
    if any(text_start.startswith(char) for char in bullet_chars):
        return True

    # Check for numbered bullets (1. 2. etc.)
    if len(text_start) >= 2 and text_start[0].isdigit() and text_start[1] in ['.', ')', ':']:
        return True

    return False
