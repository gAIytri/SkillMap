"""
Create a minimal template to test if docxtpl works
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_minimal():
    doc = Document()

    # Just test the basics
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.add_run('{{personal_info.name}}')

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run('{{personal_info.email}}')

    doc.add_paragraph()

    # Test a simple loop
    doc.add_paragraph('{% for exp in experience %}')
    doc.add_paragraph('{{exp.company}} - {{exp.title}}')
    doc.add_paragraph('{% endfor %}')

    output_path = '/Users/sidharthraj/Gaiytri projects/SkillMap/backend/templates/test_template.docx'
    doc.save(output_path)
    print(f"✅ Minimal template created: {output_path}")

if __name__ == '__main__':
    create_minimal()
