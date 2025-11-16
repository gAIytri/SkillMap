"""
Create working template with FIXED section order
We'll add dynamic reordering in Phase 4 after core functionality works
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_working_template():
    """Create a clean, working template"""
    doc = Document()

    # ==================== HEADER ====================
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.add_run('{{personal_info.name}}')
    run.font.size = Pt(18)
    run.font.bold = True

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run('{{personal_info.location}} | {{personal_info.email}} | {{personal_info.phone}}')

    links = doc.add_paragraph()
    links.alignment = WD_ALIGN_PARAGRAPH.CENTER
    links_run = links.add_run('{%if personal_info.linkedin%}{{personal_info.linkedin}}{%endif%}{%if personal_info.github%} | {{personal_info.github}}{%endif%}{%if personal_info.portfolio%} | {{personal_info.portfolio}}{%endif%}')
    links_run.font.color.rgb = RGBColor(0, 0, 255)

    doc.add_paragraph()

    # ==================== PROFESSIONAL SUMMARY ====================
    doc.add_paragraph('{%if professional_summary%}')

    summary_heading = doc.add_paragraph('PROFESSIONAL SUMMARY')
    summary_heading.runs[0].font.bold = True
    summary_heading.runs[0].font.size = Pt(11)

    summary = doc.add_paragraph('{{professional_summary}}')
    summary.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph('{%endif%}')

    # ==================== EDUCATION ====================
    doc.add_paragraph('{%if education%}')

    edu_heading = doc.add_paragraph('EDUCATION')
    edu_heading.runs[0].font.bold = True
    edu_heading.runs[0].font.size = Pt(11)

    doc.add_paragraph('{%p for edu in education%}')

    school = doc.add_paragraph()
    school_run = school.add_run('{{edu.institution}}, {{edu.location}}')
    school_run.font.bold = True

    degree = doc.add_paragraph()
    degree_run = degree.add_run('{{edu.degree}}{%if edu.gpa%} – GPA: {{edu.gpa}}{%endif%}')
    degree_run.font.italic = True

    date = doc.add_paragraph('{{edu.graduation_date}}')
    date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date.runs[0].font.bold = True

    doc.add_paragraph('{%endfor%}')
    doc.add_paragraph('{%endif%}')

    # ==================== EXPERIENCE ====================
    doc.add_paragraph('{%if experience%}')

    exp_heading = doc.add_paragraph('EXPERIENCE')
    exp_heading.runs[0].font.bold = True
    exp_heading.runs[0].font.size = Pt(11)

    doc.add_paragraph('{%p for exp in experience%}')

    company = doc.add_paragraph()
    company_run = company.add_run('{{exp.company}} – {{exp.title}}, {{exp.location}}')
    company_run.font.bold = True

    exp_date = doc.add_paragraph('{{exp.start_date}} – {{exp.end_date}}')
    exp_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    exp_date.runs[0].font.bold = True

    doc.add_paragraph('{%p for bullet in exp.bullets%}')
    doc.add_paragraph('{{bullet}}', style='List Bullet')
    doc.add_paragraph('{%endfor%}')

    doc.add_paragraph('{%endfor%}')
    doc.add_paragraph('{%endif%}')

    # ==================== PROJECTS ====================
    doc.add_paragraph('{%if projects%}')

    proj_heading = doc.add_paragraph('PROJECTS')
    proj_heading.runs[0].font.bold = True
    proj_heading.runs[0].font.size = Pt(11)

    doc.add_paragraph('{%p for project in projects%}')

    proj_name = doc.add_paragraph()
    proj_name_run = proj_name.add_run('{{project.name}}{%if project.technologies%} – {{project.technologies}}{%endif%}')
    proj_name_run.font.bold = True

    proj_date = doc.add_paragraph('{%if project.start_date%}{{project.start_date}} – {%endif%}{{project.end_date}}')
    proj_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    proj_date.runs[0].font.bold = True

    doc.add_paragraph('{%p if project.description is string%}')
    doc.add_paragraph('{{project.description}}', style='List Bullet')
    doc.add_paragraph('{%p else%}')
    doc.add_paragraph('{%p for desc in project.description%}')
    doc.add_paragraph('{{desc}}', style='List Bullet')
    doc.add_paragraph('{%p endfor%}')
    doc.add_paragraph('{%p endif%}')

    doc.add_paragraph('{%endfor%}')
    doc.add_paragraph('{%endif%}')

    # ==================== SKILLS ====================
    doc.add_paragraph('{%if skills%}')

    skills_heading = doc.add_paragraph('TECHNICAL SKILLS')
    skills_heading.runs[0].font.bold = True
    skills_heading.runs[0].font.size = Pt(11)

    doc.add_paragraph('{%p for skill_cat in skills%}')

    skill_line = doc.add_paragraph()
    cat_run = skill_line.add_run('{{skill_cat.category}}: ')
    cat_run.font.bold = True
    skill_line.add_run('{{skill_cat.skills|join(", ")}}')

    doc.add_paragraph('{%endfor%}')
    doc.add_paragraph('{%endif%}')

    # ==================== CERTIFICATIONS ====================
    doc.add_paragraph('{%if certifications%}')

    cert_heading = doc.add_paragraph('CERTIFICATIONS')
    cert_heading.runs[0].font.bold = True
    cert_heading.runs[0].font.size = Pt(11)

    doc.add_paragraph('{%p for cert in certifications%}')
    doc.add_paragraph('{{cert.name}}{%if cert.issuer%} – {{cert.issuer}}{%endif%}{%if cert.date%} ({{cert.date}}){%endif%}', style='List Bullet')
    doc.add_paragraph('{%endfor%}')
    doc.add_paragraph('{%endif%}')

    output_path = '/Users/sidharthraj/Gaiytri projects/SkillMap/backend/templates/base_resume_template.docx'
    doc.save(output_path)
    print(f"✅ Working template created: {output_path}")

if __name__ == '__main__':
    create_working_template()
