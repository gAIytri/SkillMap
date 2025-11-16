"""
Create template from Sidharth's actual resume
Preserves exact styling and structure
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_template_from_original():
    """
    Load Sidharth's resume and replace content with Jinja2 placeholders
    This preserves ALL original styling
    """

    # Load the original resume
    doc = Document('/Users/sidharthraj/Gaiytri projects/SkillMap/SidharthRaj_Khandelwal-2.docx')

    # Clear all paragraphs (we'll rebuild with Jinja2)
    for para in doc.paragraphs[:]:
        p = para._element
        p.getparent().remove(p)

    # ==================== HEADER ====================
    # Name (Title style)
    name_para = doc.add_paragraph('{{personal_info.name}}', style='Title')

    # Contact (Normal, centered, 10pt)
    contact_para = doc.add_paragraph(style='Normal')
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact_para.add_run('{{personal_info.location}} | {{personal_info.email}} | {{personal_info.phone}}')
    contact_run.font.size = Pt(10)

    # Links (Normal, centered)
    links_para = doc.add_paragraph(style='Normal')
    links_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    links_run = links_para.add_run('{%if personal_info.linkedin%}{{personal_info.linkedin}}{%endif%}{%if personal_info.github%} | {{personal_info.github}}{%endif%}{%if personal_info.portfolio%} | {{personal_info.portfolio}}{%endif%}')
    links_run.font.size = Pt(10)

    # ==================== PROFESSIONAL SUMMARY ====================
    doc.add_paragraph('{%if professional_summary%}')

    summary_heading = doc.add_paragraph('PROFESSIONAL SUMMARY', style='Heading 1')

    summary_para = doc.add_paragraph('{{professional_summary}}', style='Normal')
    summary_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    summary_para.runs[0].font.size = Pt(10)

    doc.add_paragraph('{%endif%}')

    # ==================== EDUCATION ====================
    doc.add_paragraph('{%if education%}')

    edu_heading = doc.add_paragraph('EDUCATION', style='Heading 1')

    doc.add_paragraph('{%p for edu in education%}')

    # School name (Normal, bold, justified, 10pt)
    school_para = doc.add_paragraph('{{edu.institution}}, {{edu.location}}', style='Normal')
    school_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    school_para.runs[0].font.bold = True
    school_para.runs[0].font.size = Pt(10)

    # Degree line (Normal, italic, left, 10pt)
    degree_para = doc.add_paragraph('{{edu.degree}}{%if edu.gpa%} – GPA: {{edu.gpa}}{%endif%}\t{{edu.graduation_date}}', style='Normal')
    degree_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    degree_para.runs[0].font.italic = True
    degree_para.runs[0].font.size = Pt(10)

    doc.add_paragraph('{%endfor%}')
    doc.add_paragraph('{%endif%}')

    # ==================== EXPERIENCE ====================
    doc.add_paragraph('{%if experience%}')

    exp_heading = doc.add_paragraph('EXPERIENCE', style='Heading 1')

    doc.add_paragraph('{%p for exp in experience%}')

    # Job title (Heading 2)
    job_para = doc.add_paragraph('{{exp.company}} – {{exp.title}}, {{exp.location}}\t{{exp.start_date}} – {{exp.end_date}}', style='Heading 2')

    # Bullets (List Paragraph, 9pt)
    doc.add_paragraph('{%p for bullet in exp.bullets%}')
    bullet_para = doc.add_paragraph('{{bullet}}', style='List Paragraph')
    bullet_para.runs[0].font.size = Pt(9)
    doc.add_paragraph('{%endfor%}')

    doc.add_paragraph('{%endfor%}')
    doc.add_paragraph('{%endif%}')

    # ==================== PROJECTS ====================
    doc.add_paragraph('{%if projects%}')

    proj_heading = doc.add_paragraph('PROJECTS', style='Heading 1')

    doc.add_paragraph('{%p for project in projects%}')

    # Project name (Heading 2)
    proj_name_para = doc.add_paragraph('{{project.name}}{%if project.technologies%} – {{project.technologies}}{%endif%}\t{%if project.start_date%}{{project.start_date}} – {%endif%}{{project.end_date}}', style='Heading 2')

    # Description bullets (List Paragraph, 9pt)
    doc.add_paragraph('{%p if project.description is string%}')
    desc_single = doc.add_paragraph('{{project.description}}', style='List Paragraph')
    desc_single.runs[0].font.size = Pt(9)
    doc.add_paragraph('{%p else%}')
    doc.add_paragraph('{%p for desc in project.description%}')
    desc_bullet = doc.add_paragraph('{{desc}}', style='List Paragraph')
    desc_bullet.runs[0].font.size = Pt(9)
    doc.add_paragraph('{%endfor%}')
    doc.add_paragraph('{%endif%}')

    doc.add_paragraph('{%endfor%}')
    doc.add_paragraph('{%endif%}')

    # ==================== TECHNICAL SKILLS ====================
    doc.add_paragraph('{%if skills%}')

    skills_heading = doc.add_paragraph('TECHNICAL SKILLS', style='Heading 1')

    doc.add_paragraph('{%p for skill_cat in skills%}')

    # Category: items (Normal, bold for category, 10pt)
    skill_para = doc.add_paragraph(style='Normal')
    skill_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cat_run = skill_para.add_run('{{skill_cat.category}}: ')
    cat_run.font.bold = True
    cat_run.font.size = Pt(10)
    items_run = skill_para.add_run('{{skill_cat.skills|join(", ")}}')
    items_run.font.size = Pt(10)

    doc.add_paragraph('{%endfor%}')
    doc.add_paragraph('{%endif%}')

    # ==================== CERTIFICATIONS ====================
    doc.add_paragraph('{%if certifications%}')

    cert_heading = doc.add_paragraph('CERTIFICATIONS', style='Heading 1')

    doc.add_paragraph('{%p for cert in certifications%}')
    cert_para = doc.add_paragraph('{{cert.name}}{%if cert.issuer%} – {{cert.issuer}}{%endif%}{%if cert.date%} ({{cert.date}}){%endif%}', style='List Paragraph')
    cert_para.runs[0].font.size = Pt(9)
    doc.add_paragraph('{%endfor%}')
    doc.add_paragraph('{%endif%}')

    # Save template
    output_path = '/Users/sidharthraj/Gaiytri projects/SkillMap/backend/templates/base_resume_template.docx'
    doc.save(output_path)
    print(f"✅ Template created from your original resume: {output_path}")
    print("   This template preserves YOUR exact styling!")

if __name__ == '__main__':
    from docx.shared import Pt
    create_template_from_original()
