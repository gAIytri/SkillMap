"""
Script to create the base resume template with Jinja2 placeholders
This creates a DOCX template matching Sidharth's resume structure
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_template():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # ==================== HEADER ====================
    # Name
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run('{{ personal_info.name }}')
    name_run.font.size = Pt(20)
    name_run.font.bold = True
    name_run.font.name = 'Calibri'

    # Contact Info
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_text = '{{ personal_info.location }} | {{ personal_info.email }} | {{ personal_info.phone }}'
    contact_run = contact_para.add_run(contact_text)
    contact_run.font.size = Pt(11)

    # Links
    links_para = doc.add_paragraph()
    links_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    links_text = '{% if personal_info.linkedin %}{{ personal_info.linkedin }}{% endif %}{% if personal_info.github %} | {{ personal_info.github }}{% endif %}{% if personal_info.portfolio %} | {{ personal_info.portfolio }}{% endif %}'
    links_run = links_para.add_run(links_text)
    links_run.font.size = Pt(11)
    links_run.font.color.rgb = RGBColor(0, 0, 255)  # Blue for links

    # ==================== SECTION TEMPLATE ====================
    # We'll use a loop to render sections in custom order

    # Add Jinja2 logic for dynamic section ordering
    doc.add_paragraph()  # Spacing

    # Loop through sections in user's custom order
    section_loop = doc.add_paragraph()
    section_loop.add_run('{% for section in sections %}')

    # Professional Summary
    summary_check = doc.add_paragraph()
    summary_check.add_run("{% if section == 'professional_summary' and professional_summary %}")

    summary_heading = doc.add_paragraph('PROFESSIONAL SUMMARY')
    summary_heading.runs[0].font.bold = True
    summary_heading.runs[0].font.size = Pt(12)
    # Add underline using XML
    summary_heading._element.get_or_add_pPr().append(
        parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:bottom w:val="single" w:sz="6" w:space="1" w:color="000000"/></w:pBdr>')
    )

    summary_content = doc.add_paragraph('{{ professional_summary }}')
    summary_content.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    end_summary = doc.add_paragraph()
    end_summary.add_run('{% endif %}')

    # Education
    edu_check = doc.add_paragraph()
    edu_check.add_run("{% if section == 'education' and education %}")

    edu_heading = doc.add_paragraph('EDUCATION')
    edu_heading.runs[0].font.bold = True
    edu_heading.runs[0].font.size = Pt(12)
    edu_heading._element.get_or_add_pPr().append(
        parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:bottom w:val="single" w:sz="6" w:space="1" w:color="000000"/></w:pBdr>')
    )

    edu_loop = doc.add_paragraph()
    edu_loop.add_run('{% for edu in education %}')

    # School name and dates
    school_para = doc.add_paragraph()
    school_run = school_para.add_run('{{ edu.institution }}, {{ edu.location }}')
    school_run.font.bold = True

    # Degree and GPA
    degree_para = doc.add_paragraph()
    degree_run = degree_para.add_run('{{ edu.degree }}')
    degree_run.font.italic = True
    if_gpa = degree_para.add_run(' – GPA: {{ edu.gpa }}')
    if_gpa.font.italic = True

    # Date (right-aligned)
    date_para = doc.add_paragraph('{{ edu.graduation_date }}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_para.runs[0].font.bold = True

    end_edu_loop = doc.add_paragraph()
    end_edu_loop.add_run('{% endfor %}')

    end_edu = doc.add_paragraph()
    end_edu.add_run('{% endif %}')

    # Experience
    exp_check = doc.add_paragraph()
    exp_check.add_run("{% if section == 'experience' and experience %}")

    exp_heading = doc.add_paragraph('EXPERIENCE')
    exp_heading.runs[0].font.bold = True
    exp_heading.runs[0].font.size = Pt(12)
    exp_heading._element.get_or_add_pPr().append(
        parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:bottom w:val="single" w:sz="6" w:space="1" w:color="000000"/></w:pBdr>')
    )

    exp_loop = doc.add_paragraph()
    exp_loop.add_run('{% for exp in experience %}')

    # Company and title
    company_para = doc.add_paragraph()
    company_run = company_para.add_run('{{ exp.company }} – {{ exp.title }}, {{ exp.location }}')
    company_run.font.bold = True

    # Dates (right-aligned)
    exp_date_para = doc.add_paragraph('{{ exp.start_date }} – {{ exp.end_date }}')
    exp_date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    exp_date_para.runs[0].font.bold = True

    # Bullets
    bullet_loop = doc.add_paragraph()
    bullet_loop.add_run('{% for bullet in exp.bullets %}')

    bullet_para = doc.add_paragraph('{{ bullet }}', style='List Bullet')

    end_bullet_loop = doc.add_paragraph()
    end_bullet_loop.add_run('{% endfor %}')

    end_exp_loop = doc.add_paragraph()
    end_exp_loop.add_run('{% endfor %}')

    end_exp = doc.add_paragraph()
    end_exp.add_run('{% endif %}')

    # Projects
    proj_check = doc.add_paragraph()
    proj_check.add_run("{% if section == 'projects' and projects %}")

    proj_heading = doc.add_paragraph('PROJECTS')
    proj_heading.runs[0].font.bold = True
    proj_heading.runs[0].font.size = Pt(12)
    proj_heading._element.get_or_add_pPr().append(
        parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:bottom w:val="single" w:sz="6" w:space="1" w:color="000000"/></w:pBdr>')
    )

    proj_loop = doc.add_paragraph()
    proj_loop.add_run('{% for project in projects %}')

    # Project name and tech stack
    proj_name_para = doc.add_paragraph()
    proj_name_run = proj_name_para.add_run('{{ project.name }}')
    proj_name_run.font.bold = True
    if 'technologies' in '{{ project }}':
        tech_run = proj_name_para.add_run(' – {{ project.technologies }}')
        tech_run.font.bold = False

    # Dates (right-aligned)
    proj_date_para = doc.add_paragraph('{% if project.start_date %}{{ project.start_date }} – {% endif %}{{ project.end_date }}')
    proj_date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    proj_date_para.runs[0].font.bold = True

    # Project description bullets
    proj_bullet_loop = doc.add_paragraph()
    proj_bullet_loop.add_run('{% if project.description is string %}')

    desc_para = doc.add_paragraph('{{ project.description }}', style='List Bullet')

    else_check = doc.add_paragraph()
    else_check.add_run('{% else %}')

    desc_loop = doc.add_paragraph()
    desc_loop.add_run('{% for desc in project.description %}')

    desc_bullet_para = doc.add_paragraph('{{ desc }}', style='List Bullet')

    end_desc_loop = doc.add_paragraph()
    end_desc_loop.add_run('{% endfor %}')

    end_desc_check = doc.add_paragraph()
    end_desc_check.add_run('{% endif %}')

    end_proj_loop = doc.add_paragraph()
    end_proj_loop.add_run('{% endfor %}')

    end_proj = doc.add_paragraph()
    end_proj.add_run('{% endif %}')

    # Skills
    skills_check = doc.add_paragraph()
    skills_check.add_run("{% if section == 'skills' and skills %}")

    skills_heading = doc.add_paragraph('TECHNICAL SKILLS')
    skills_heading.runs[0].font.bold = True
    skills_heading.runs[0].font.size = Pt(12)
    skills_heading._element.get_or_add_pPr().append(
        parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:bottom w:val="single" w:sz="6" w:space="1" w:color="000000"/></w:pBdr>')
    )

    skills_loop = doc.add_paragraph()
    skills_loop.add_run('{% for skill_cat in skills %}')

    # Category and items
    skill_para = doc.add_paragraph()
    cat_run = skill_para.add_run('{{ skill_cat.category }}: ')
    cat_run.font.bold = True
    skill_para.add_run('{{ skill_cat.items|join(", ") }}')

    end_skills_loop = doc.add_paragraph()
    end_skills_loop.add_run('{% endfor %}')

    end_skills = doc.add_paragraph()
    end_skills.add_run('{% endif %}')

    # Certifications
    cert_check = doc.add_paragraph()
    cert_check.add_run("{% if section == 'certifications' and certifications %}")

    cert_heading = doc.add_paragraph('CERTIFICATIONS')
    cert_heading.runs[0].font.bold = True
    cert_heading.runs[0].font.size = Pt(12)
    cert_heading._element.get_or_add_pPr().append(
        parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:bottom w:val="single" w:sz="6" w:space="1" w:color="000000"/></w:pBdr>')
    )

    cert_loop = doc.add_paragraph()
    cert_loop.add_run('{% for cert in certifications %}')

    cert_para = doc.add_paragraph('{{ cert.name }}{% if cert.issuer %} – {{ cert.issuer }}{% endif %}{% if cert.date %} ({{ cert.date }}){% endif %}', style='List Bullet')

    end_cert_loop = doc.add_paragraph()
    end_cert_loop.add_run('{% endfor %}')

    end_cert = doc.add_paragraph()
    end_cert.add_run('{% endif %}')

    # End section loop
    end_section_loop = doc.add_paragraph()
    end_section_loop.add_run('{% endfor %}')

    # Save template
    doc.save('/Users/sidharthraj/Gaiytri projects/SkillMap/backend/templates/base_resume_template.docx')
    print("✅ Template created successfully at: backend/templates/base_resume_template.docx")

if __name__ == '__main__':
    create_template()
