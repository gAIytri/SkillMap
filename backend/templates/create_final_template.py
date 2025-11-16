"""
Create final template WITHOUT section ordering loops
Section ordering will be handled in Python, not in the template
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_final_template():
    """Create template with all sections - ordering handled by Python"""
    doc = Document()

    # ==================== HEADER ====================
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.add_run('{{personal_info.name}}')
    run.font.size = Pt(18)
    run.font.bold = True

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run('{{personal_info.location}} | {{personal_info.email}} | {{personal_info.phone}}')

    links = doc.add_paragraph()
    links.alignment = WD_ALIGN_PARAGRAPH.CENTER
    links_run = links.add_run('{%if personal_info.linkedin%}{{personal_info.linkedin}}{%endif%}{%if personal_info.github%} | {{personal_info.github}}{%endif%}{%if personal_info.portfolio%} | {{personal_info.portfolio}}{%endif%}')
    links_run.font.color.rgb = RGBColor(0, 0, 255)

    doc.add_paragraph()

    # ==================== RENDERED_SECTIONS ====================
    # This will be populated by Python before rendering
    doc.add_paragraph('{{rendered_sections}}')

    output_path = '/Users/sidharthraj/Gaiytri projects/SkillMap/backend/templates/base_resume_template.docx'
    doc.save(output_path)
    print(f"✅ Final template created: {output_path}")

if __name__ == '__main__':
    create_final_template()
