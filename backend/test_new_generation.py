"""
Test the new programmatic DOCX generation service
"""

from services.resume_extractor import extract_resume
from services.docx_generation_service import generate_resume_from_json, get_default_section_order
from docx import Document
import json

print("="*80)
print("TESTING NEW PROGRAMMATIC DOCX GENERATION")
print("="*80)

# Step 1: Extract resume from your DOCX
print("\nStep 1: Extracting resume...")
with open('/Users/sidharthraj/Gaiytri projects/SkillMap/SidharthRaj_Khandelwal-2.docx', 'rb') as f:
    docx_bytes = f.read()

resume_json = extract_resume(docx_bytes)

print("✅ Extraction complete!")
print(f"\nExtracted sections:")
print(f"  - Name: {resume_json.get('personal_info', {}).get('name', 'N/A')}")
print(f"  - Education entries: {len(resume_json.get('education', []))}")
print(f"  - Experience entries: {len(resume_json.get('experience', []))}")
print(f"  - Projects: {len(resume_json.get('projects', []))}")
print(f"  - Skill categories: {len(resume_json.get('skills', []))}")

# Step 2: Generate resume from JSON
print("\n" + "="*80)
print("Step 2: Generating resume programmatically...")

section_order = get_default_section_order()
output_bytes = generate_resume_from_json(
    resume_json=resume_json,
    base_resume_docx=docx_bytes,  # Use original as style reference
    section_order=section_order
)

print(f"✅ Generation complete! Size: {len(output_bytes)} bytes")

# Step 3: Save to test file
output_path = '/Users/sidharthraj/Gaiytri projects/SkillMap/backend/test_generated_resume.docx'
with open(output_path, 'wb') as f:
    f.write(output_bytes)

print(f"✅ Saved to: {output_path}")

# Step 4: Analyze generated resume
print("\n" + "="*80)
print("Step 3: Analyzing generated resume...")

doc = Document(output_path)

print(f"\nTotal paragraphs: {len(doc.paragraphs)}")

print("\nFirst 30 paragraphs:")
count = 0
for idx, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if not text:
        continue

    count += 1
    if count > 30:
        break

    style = para.style.name if para.style else "None"

    # Get font info
    font_info = ""
    if para.runs:
        run = para.runs[0]
        font_size = f"{run.font.size.pt}pt" if run.font.size else "Default"
        is_bold = "Bold" if run.font.bold else ""
        is_italic = "Italic" if run.font.italic else ""
        is_underline = "Underline" if run.font.underline else ""
        font_info = f"{font_size} {is_bold} {is_italic} {is_underline}".strip()

    print(f"\n[{count}] Style: {style} | Font: {font_info}")
    print(f"    {text[:100]}")

print("\n" + "="*80)
print("SUMMARY")
print("="*80)
print("✅ Extraction: Working")
print("✅ Generation: Working")
print("✅ Styles preserved: Check the DOCX file")
print(f"\nNext: Open {output_path} and compare with your original resume!")
